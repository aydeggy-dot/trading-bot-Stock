"""
NGX Trading Bot - Cron Jobs Reference Document Generator
Generates a comprehensive .docx documenting all 24 scheduled tasks across 14 classes.

Usage:  python docs/generate_cron_docs.py
Output: docs/word/CRON_JOBS_REFERENCE.docx
"""
import base64
import io
import time
from pathlib import Path

import requests
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ═══════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════════
OUTPUT_DIR = Path(__file__).parent / "word"
BRAND_BLUE = RGBColor(0x33, 0x9A, 0xF0)
BRAND_GREEN = RGBColor(0x51, 0xCF, 0x66)
BRAND_RED = RGBColor(0xF0, 0x3E, 0x3E)
BRAND_ORANGE = RGBColor(0xFF, 0xA9, 0x4D)
BRAND_DARK = RGBColor(0x21, 0x25, 0x29)
BRAND_GRAY = RGBColor(0x86, 0x8E, 0x96)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ═══════════════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS (adapted from generate_docs.py)
# ═══════════════════════════════════════════════════════════════════════════════

def render_mermaid(mermaid_code: str, max_retries: int = 2) -> bytes | None:
    encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("utf-8")
    url = f"https://mermaid.ink/img/{encoded}?type=png&bgColor=white"
    for attempt in range(max_retries + 1):
        try:
            resp = requests.get(url, timeout=30)
            if resp.status_code == 200 and len(resp.content) > 100:
                return resp.content
        except requests.RequestException:
            pass
        if attempt < max_retries:
            time.sleep(2)
    print(f"  [WARNING] Mermaid render failed after {max_retries + 1} attempts")
    return None


def setup_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BRAND_DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for level, (size, bold, color) in {
        "Heading 1": (Pt(24), True, BRAND_BLUE),
        "Heading 2": (Pt(18), True, BRAND_DARK),
        "Heading 3": (Pt(14), True, BRAND_BLUE),
        "Heading 4": (Pt(12), True, BRAND_DARK),
    }.items():
        h = doc.styles[level]
        h.font.name = "Calibri"
        h.font.size = size
        h.font.bold = bold
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(18) if level != "Heading 1" else Pt(6)
        h.paragraph_format.space_after = Pt(8)


def add_title_page(doc: Document, title: str, subtitle: str, audience: str):
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = BRAND_BLUE
    run.font.name = "Calibri"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_GRAY
    run.font.name = "Calibri"
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Audience: {audience}")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = BRAND_DARK
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NGX Trading Bot")
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_BLUE
    run.font.bold = True
    doc.add_page_break()


def add_styled_table(doc: Document, headers: list[str], rows: list[list[str]],
                     col_widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="339AF0"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F3F5"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph("")
    return table


def add_diagram(doc: Document, mermaid_code: str, caption: str = "",
                width: float = 6.0):
    img_data = render_mermaid(mermaid_code)
    if img_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(img_data), width=Inches(width))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = BRAND_GRAY
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Diagram: {caption}]")
        run.font.italic = True
        run.font.color.rgb = BRAND_RED


def add_callout(doc: Document, text: str, color: RGBColor = BRAND_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = color
    run.font.italic = True
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
    p._p.get_or_add_pPr().append(shading)


def add_bullet(doc: Document, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        if not p.runs:
            run = p.add_run(text)
            run.font.size = Pt(11)
        else:
            p.text = text


def add_section_divider(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 60)
    run.font.color.rgb = RGBColor(0xDE, 0xE2, 0xE6)
    run.font.size = Pt(8)


def save_doc(doc: Document, filename: str):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc.save(str(path))
    print(f"  Saved: {path}")


# ─── New Helpers ──────────────────────────────────────────────────────────────

def add_page_number(paragraph):
    """Insert a PAGE field code for auto page numbering."""
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld2)


def apply_header_footer(section, title="Cron Jobs Reference"):
    """Apply header text and footer page numbers to a section."""
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(title)
    run.font.size = Pt(8)
    run.font.color.rgb = BRAND_GRAY
    run.font.name = "Calibri"
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


def switch_to_landscape(doc):
    """Add a landscape section and return it."""
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    apply_header_footer(section)
    return section


def switch_to_portrait(doc):
    """Return to portrait orientation."""
    section = doc.add_section()
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    apply_header_footer(section)
    return section


def add_toc_field(doc):
    """Insert a Word TOC field code (update in Word with Ctrl+A, F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld1)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    run._r.append(fld2)
    run2 = p.add_run("Right-click and select Update Field to generate Table of Contents")
    run2.font.color.rgb = BRAND_GRAY
    run2.font.italic = True
    run2.font.size = Pt(10)
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    run2._r.append(fld3)


# ═══════════════════════════════════════════════════════════════════════════════
# DATA REGISTRY — Single source of truth for all 24 scheduled jobs
# ═══════════════════════════════════════════════════════════════════════════════

SCHEDULERS = [
    # ── 1. DataCollectionScheduler ────────────────────────────────────────────
    dict(
        name="DataCollectionScheduler",
        display="Data Collection Scheduler",
        path="src/main/java/com/ngxbot/data/scheduler/DataCollectionScheduler.java",
        category="Data Collection", module="data",
        overview="Foundation of the data pipeline. Collects OHLCV price bars from "
                 "EODHD API, scrapes NGX website as backup, fetches ETF NAVs, and "
                 "scans RSS feeds for financial news.",
        deps=[
            ("Client", "EodhdApiClient", "EODHD REST API for OHLCV data"),
            ("Scraper", "NgxWebScraper", "NGX website backup price scraper"),
            ("Scraper", "EtfNavScraper", "ETF NAV scraper from fund managers"),
            ("Scraper", "NewsRssScraper", "Financial news RSS feed parser"),
            ("Config", "TradingProperties", "Watchlist symbols and config"),
        ],
        tables_r=[], tables_w=["ohlcv_bar", "news_item", "market_index", "etf_valuation"],
        kill_switch=False,
        jobs=[
            dict(method="preMarketNewsScan", cron="0 0 9 * * MON-FRI",
                 time="09:00", days="Mon-Fri",
                 purpose="Scans financial RSS feeds for pre-market news"),
            dict(method="pullEodhdData", cron="0 0 15 * * MON-FRI",
                 time="15:00", days="Mon-Fri",
                 purpose="Fetches 5-day OHLCV bars from EODHD for all watchlist symbols"),
            dict(method="scrapeNgxBackup", cron="0 5 15 * * MON-FRI",
                 time="15:05", days="Mon-Fri",
                 purpose="Backup: scrapes NGX website for daily prices and ASI index"),
            dict(method="scrapeEtfNavs", cron="0 0 16 * * MON-FRI",
                 time="16:00", days="Mon-Fri",
                 purpose="Scrapes ETF Net Asset Values from fund manager sites"),
        ],
    ),
    # ── 2. NewsScraperScheduler ───────────────────────────────────────────────
    dict(
        name="NewsScraperScheduler",
        display="News Scraper Scheduler",
        path="src/main/java/com/ngxbot/news/scraper/NewsScraperScheduler.java",
        category="Data Collection", module="news",
        overview="Runs 5 web scrapers (Nairametrics, BusinessDay, Reuters RSS, "
                 "SeekingAlpha, CBN Press) three times daily, plus NGX bulletin "
                 "parsing in the after-market scan.",
        deps=[
            ("Scraper", "NairametricsScraper", "Nairametrics news scraper"),
            ("Scraper", "BusinessDayScraper", "BusinessDay news scraper"),
            ("Scraper", "ReutersRssScraper", "Reuters RSS feed scraper"),
            ("Scraper", "SeekingAlphaScraper", "SeekingAlpha news scraper"),
            ("Scraper", "CbnPressScraper", "CBN press release scraper"),
            ("Parser", "NgxBulletinParser", "NGX daily bulletin parser"),
        ],
        tables_r=[], tables_w=["news_item"],
        kill_switch=False,
        jobs=[
            dict(method="earlyMorningNewsScan", cron="0 0 8 * * MON-FRI",
                 time="08:00", days="Mon-Fri",
                 purpose="Pre-market news scan via all 5 web scrapers"),
            dict(method="midDayNewsScan", cron="0 30 12 * * MON-FRI",
                 time="12:30", days="Mon-Fri",
                 purpose="Mid-day news refresh via all 5 web scrapers"),
            dict(method="afterMarketNewsScan", cron="0 0 17 * * MON-FRI",
                 time="17:00", days="Mon-Fri",
                 purpose="After-market scan: 5 web scrapers + NGX bulletin parser"),
        ],
    ),
    # ── 3. AiAnalysisScheduler ────────────────────────────────────────────────
    dict(
        name="AiAnalysisScheduler",
        display="AI Analysis Scheduler",
        path="src/main/java/com/ngxbot/ai/scheduler/AiAnalysisScheduler.java",
        category="Analysis", module="ai",
        overview="Processes unanalyzed news articles via Claude API in batches "
                 "(every 15 min), and runs cross-article sentiment synthesis daily.",
        deps=[
            ("Service", "AiNewsAnalyzer", "Claude API article analyzer"),
            ("Service", "AiCrossArticleSynthesizer", "Multi-article synthesis"),
            ("Repo", "NewsItemRepository", "Unprocessed news query"),
            ("Repo", "AiAnalysisRepository", "AI results persistence"),
            ("Service", "AiCostTracker", "API spending budget tracker"),
            ("Config", "AiProperties", "AI module config"),
        ],
        tables_r=["news_item", "ai_analysis", "ai_cost_ledger"],
        tables_w=["news_item", "ai_analysis"],
        kill_switch=False,
        jobs=[
            dict(method="processUnanalyzedArticles", cron="fixedDelay=900000",
                 time="Every 15m", days="Always",
                 purpose="Batch-analyzes unprocessed news via Claude API"),
            dict(method="crossArticleSynthesis", cron="0 0 21 * * MON-FRI",
                 time="21:00", days="Mon-Fri",
                 purpose="Synthesizes consensus from multiple articles per symbol"),
        ],
    ),
    # ── 4. SignalGenerationScheduler ──────────────────────────────────────────
    dict(
        name="SignalGenerationScheduler",
        display="Signal Generation Scheduler",
        path="src/main/java/com/ngxbot/signal/SignalGenerationScheduler.java",
        category="Signal", module="signal",
        overview="Generates trade signals by running composite scoring across all "
                 "strategies. NGX signals at 15:15 (after EODHD pull), US at 21:30.",
        deps=[
            ("Service", "CompositeSignalScorer", "Multi-strategy signal scorer"),
            ("Repo", "TradeSignalRepository", "Signal persistence"),
            ("Util", "ObjectMapper", "JSON serialization"),
        ],
        tables_r=["ohlcv_bar", "position", "ai_analysis"],
        tables_w=["trade_signal"],
        kill_switch=False,
        jobs=[
            dict(method="generateNgxSignals", cron="0 15 15 * * MON-FRI",
                 time="15:15", days="Mon-Fri",
                 purpose="Generates NGX trade signals from all non-US strategies"),
            dict(method="generateUsSignals", cron="0 30 21 * * MON-FRI",
                 time="21:30", days="Mon-Fri",
                 purpose="Generates US trade signals from US+dual-market strategies"),
        ],
    ),
    # ── 5. DiscoveryScheduler ─────────────────────────────────────────────────
    dict(
        name="DiscoveryScheduler",
        display="Discovery Scheduler",
        path="src/main/java/com/ngxbot/discovery/scheduler/DiscoveryScheduler.java",
        category="Discovery", module="discovery",
        overview="Manages the stock discovery pipeline: weekly EODHD screening, "
                 "daily observation evaluation, and daily demotion checks.",
        deps=[
            ("Client", "EodhdScreenerClient", "EODHD stock screener"),
            ("Service", "CandidateEvaluator", "Candidate scoring"),
            ("Policy", "PromotionPolicy", "OBSERVATION to PROMOTED rules"),
            ("Policy", "DemotionPolicy", "PROMOTED to DEMOTED rules"),
            ("Repo", "DiscoveredStockRepository", "Discovery persistence"),
            ("Config", "DiscoveryProperties", "Module config"),
            ("Service", "WatchlistManager", "Watchlist updates"),
        ],
        tables_r=["discovered_stock"], tables_w=["discovered_stock", "discovery_event"],
        kill_switch=False,
        jobs=[
            dict(method="weeklyScreenerScan", cron="0 0 6 * * SUN",
                 time="06:00", days="Sunday",
                 purpose="Weekly EODHD screener scan for new NGX candidates"),
            dict(method="dailyObservationCheck", cron="0 0 20 * * MON-FRI",
                 time="20:00", days="Mon-Fri",
                 purpose="Evaluates OBSERVATION stocks for promotion criteria"),
            dict(method="dailyDemotionCheck", cron="0 30 20 * * MON-FRI",
                 time="20:30", days="Mon-Fri",
                 purpose="Checks PROMOTED stocks for demotion conditions"),
        ],
    ),
    # ── 6. DcaScheduler ──────────────────────────────────────────────────────
    dict(
        name="DcaScheduler",
        display="DCA Scheduler",
        path="src/main/java/com/ngxbot/longterm/scheduler/DcaScheduler.java",
        category="Long-term", module="longterm",
        overview="Dollar-cost averaging executor. Runs daily at 10:15 AM, checks "
                 "if today is a configured DCA day, and executes for NGX/US markets.",
        deps=[
            ("Service", "DcaExecutor", "DCA order execution"),
            ("Config", "LongtermProperties", "DCA config (days, amounts)"),
        ],
        tables_r=[], tables_w=[],
        kill_switch=False,
        jobs=[
            dict(method="checkAndExecuteDca", cron="0 15 10 * * *",
                 time="10:15", days="Daily",
                 purpose="Checks if today is DCA day, executes for NGX and/or US"),
        ],
    ),
    # ── 7. DividendCheckScheduler ─────────────────────────────────────────────
    dict(
        name="DividendCheckScheduler",
        display="Dividend Check Scheduler",
        path="src/main/java/com/ngxbot/longterm/scheduler/DividendCheckScheduler.java",
        category="Long-term", module="longterm",
        overview="Weekly Monday check for upcoming ex-dividend dates within alert "
                 "window, and processes pending dividend reinvestments.",
        deps=[
            ("Service", "DividendTracker", "Ex-date tracking"),
            ("Service", "DividendReinvestmentService", "Reinvestment execution"),
            ("Config", "LongtermProperties", "Dividend config"),
        ],
        tables_r=["dividend_event"], tables_w=["dividend_event"],
        kill_switch=False,
        jobs=[
            dict(method="weeklyDividendCheck", cron="0 0 9 * * MON",
                 time="09:00", days="Monday",
                 purpose="Checks upcoming ex-dates and processes reinvestments"),
        ],
    ),
    # ── 8. FundamentalUpdateScheduler ─────────────────────────────────────────
    dict(
        name="FundamentalUpdateScheduler",
        display="Fundamental Update Scheduler",
        path="src/main/java/com/ngxbot/longterm/scheduler/FundamentalUpdateScheduler.java",
        category="Long-term", module="longterm",
        overview="Weekly Sunday refresh of fundamental scores for all core holdings "
                 "(NGX + US) via the FundamentalScreener.",
        deps=[
            ("Service", "FundamentalScreener", "Fundamental scoring engine"),
            ("Service", "CorePortfolioManager", "Core holdings manager"),
        ],
        tables_r=["core_holding"], tables_w=["core_holding"],
        kill_switch=False,
        jobs=[
            dict(method="weeklyFundamentalUpdate", cron="0 0 20 * * SUN",
                 time="20:00", days="Sunday",
                 purpose="Refreshes fundamental scores for all core holdings"),
        ],
    ),
    # ── 9. RebalanceScheduler ─────────────────────────────────────────────────
    dict(
        name="RebalanceScheduler",
        display="Rebalance Scheduler",
        path="src/main/java/com/ngxbot/longterm/scheduler/RebalanceScheduler.java",
        category="Long-term", module="longterm",
        overview="Quarterly rebalance on 1st of Jan/Apr/Jul/Oct at 15:00. Refreshes "
                 "market values, calculates drift, generates rebalance actions.",
        deps=[
            ("Service", "PortfolioRebalancer", "Rebalance action generator"),
            ("Service", "CorePortfolioManager", "Market value refresh"),
            ("Config", "LongtermProperties", "Rebalance config"),
        ],
        tables_r=["core_holding"], tables_w=["core_holding", "rebalance_action"],
        kill_switch=False,
        jobs=[
            dict(method="quarterlyRebalanceCheck", cron="0 0 15 1 1,4,7,10 *",
                 time="15:00", days="1st of Jan/Apr/Jul/Oct",
                 purpose="Quarterly portfolio rebalance if drift exceeds threshold"),
        ],
    ),
    # ── 10. CircuitBreaker ────────────────────────────────────────────────────
    dict(
        name="CircuitBreaker",
        display="Circuit Breaker",
        path="src/main/java/com/ngxbot/risk/service/CircuitBreaker.java",
        category="Risk", module="risk",
        overview="Monitors portfolio P&L every 10 minutes during combined trading "
                 "hours. Halts SATELLITE trading on excessive losses. CORE pool unaffected.",
        deps=[
            ("Config", "RiskProperties", "Loss thresholds"),
            ("Repo", "PortfolioSnapshotRepository", "P&L snapshot data"),
        ],
        tables_r=["portfolio_snapshot"], tables_w=[],
        kill_switch=False,
        jobs=[
            dict(method="periodicCircuitCheck", cron="0 */10 10-22 * * MON-FRI",
                 time="Every 10m (10:00-22:00)", days="Mon-Fri",
                 purpose="Checks daily/weekly/cross-market loss thresholds"),
        ],
    ),
    # ── 11. SettlementCashTracker ─────────────────────────────────────────────
    dict(
        name="SettlementCashTracker",
        display="Settlement Cash Tracker",
        path="src/main/java/com/ngxbot/risk/service/SettlementCashTracker.java",
        category="Risk", module="risk",
        overview="Daily settlement processing at 9:00 AM. Moves settled amounts "
                 "from settling pool to available cash (T+2 NGX, T+1 US). In-memory ledger.",
        deps=[],
        tables_r=[], tables_w=[],
        kill_switch=False,
        jobs=[
            dict(method="dailySettlementProcessing", cron="0 0 9 * * MON-FRI",
                 time="09:00", days="Mon-Fri",
                 purpose="Processes T+2/T+1 settlements, moves to available cash"),
        ],
    ),
    # ── 12. StopLossMonitor ───────────────────────────────────────────────────
    dict(
        name="StopLossMonitor",
        display="Stop-Loss Monitor",
        path="src/main/java/com/ngxbot/risk/service/StopLossMonitor.java",
        category="Risk", module="risk",
        overview="Monitors open positions every 5 minutes during market hours. "
                 "NGX positions during 10:00-14:00, US positions during 15:00-21:00.",
        deps=[
            ("Repo", "PositionRepository", "Open position queries"),
            ("Repo", "OhlcvRepository", "Latest price data"),
        ],
        tables_r=["position", "ohlcv_bar"], tables_w=["position"],
        kill_switch=False,
        jobs=[
            dict(method="monitorNgxStops", cron="0 */5 10-14 * * MON-FRI",
                 time="Every 5m (10:00-14:00)", days="Mon-Fri",
                 purpose="Checks NGX positions (.XNSA) against stop-loss levels"),
            dict(method="monitorUsStops", cron="0 */5 15-21 * * MON-FRI",
                 time="Every 5m (15:00-21:00)", days="Mon-Fri",
                 purpose="Checks US positions against stop-loss levels"),
        ],
    ),
    # ── 13. PortfolioReconciler ───────────────────────────────────────────────
    dict(
        name="PortfolioReconciler",
        display="Portfolio Reconciler",
        path="src/main/java/com/ngxbot/execution/service/PortfolioReconciler.java",
        category="Execution", module="execution",
        overview="Pre-market position reconciliation at 9:30 AM. Compares DB positions "
                 "against broker holdings. Activates kill switch on mismatch.",
        deps=[
            ("Repo", "PositionRepository", "DB position data"),
            ("Service", "KillSwitchService", "Kill switch activation"),
            ("Service", "NotificationRouter", "Urgent alert routing"),
            ("Gateway", "BrokerGateway", "Broker portfolio query"),
        ],
        tables_r=["position"], tables_w=[],
        kill_switch=True,
        jobs=[
            dict(method="dailyReconciliation", cron="0 30 9 * * MON-FRI",
                 time="09:30", days="Mon-Fri",
                 purpose="Reconciles DB positions vs broker holdings; kill switch on mismatch"),
        ],
    ),
    # ── 14. CashReconciler ────────────────────────────────────────────────────
    dict(
        name="CashReconciler",
        display="Cash Reconciler",
        path="src/main/java/com/ngxbot/execution/service/CashReconciler.java",
        category="Execution", module="execution",
        overview="Pre-market cash reconciliation at 9:35 AM. Compares internal cash "
                 "ledger against broker's actual balance. Kill switch on >2% mismatch.",
        deps=[
            ("Service", "SettlementCashTracker", "Internal cash ledger"),
            ("Service", "KillSwitchService", "Kill switch activation"),
            ("Service", "NotificationRouter", "Urgent alert routing"),
            ("Gateway", "BrokerGateway", "Broker cash balance query"),
        ],
        tables_r=[], tables_w=[],
        kill_switch=True,
        jobs=[
            dict(method="dailyCashReconciliation", cron="0 35 9 * * MON-FRI",
                 time="09:35", days="Mon-Fri",
                 purpose="Reconciles internal cash vs broker balance; kill switch on mismatch"),
        ],
    ),
]

# ═══════════════════════════════════════════════════════════════════════════════
# FLOWCHARTS — One per scheduler (primary job)
# ═══════════════════════════════════════════════════════════════════════════════

FLOWCHARTS = {
    "DataCollectionScheduler": """graph TD
    A["pullEodhdData - 15:00 MON-FRI"] --> B["Get all watchlist symbols"]
    B --> C{"For each symbol"}
    C --> D["Fetch 5-day OHLCV from EODHD"]
    D --> E{"API Success?"}
    E -->|Yes| F["Persist bars to DB"]
    E -->|No| G["Log error, increment fail"]
    F --> C
    G --> C
    C -->|All done| H["Log summary: success/fail"]""",

    "NewsScraperScheduler": """graph TD
    A["afterMarketNewsScan - 17:00"] --> B["Run 5 web scrapers"]
    B --> C["Nairametrics"]
    B --> D["BusinessDay"]
    B --> E["Reuters RSS"]
    B --> F["SeekingAlpha"]
    B --> G["CBN Press"]
    C --> H["Collect all NewsItem results"]
    D --> H
    E --> H
    F --> H
    G --> H
    H --> I["Run NGX Bulletin Parser"]
    I --> J["Log total items"]""",

    "AiAnalysisScheduler": """graph TD
    A["processUnanalyzedArticles - every 15m"] --> B{"AI enabled?"}
    B -->|No| C["Skip"]
    B -->|Yes| D["Query unprocessed articles"]
    D --> E{"Any found?"}
    E -->|No| F["Skip - nothing to process"]
    E -->|Yes| G{"Budget available?"}
    G -->|No| H["Log budget exceeded"]
    G -->|Yes| I["Analyze via Claude API"]
    I --> J["Persist AI analysis"]
    J --> K["Mark articles processed"]""",

    "SignalGenerationScheduler": """graph TD
    A["generateNgxSignals - 15:15 MON-FRI"] --> B["Filter NGX strategies"]
    B --> C["Run CompositeSignalScorer"]
    C --> D["Score and rank signals"]
    D --> E["Persist TradeSignal entities"]
    E --> F["Log signal count"]""",

    "DiscoveryScheduler": """graph TD
    A["weeklyScreenerScan - SUN 06:00"] --> B{"Discovery enabled?"}
    B -->|No| C["Skip"]
    B -->|Yes| D["Screen NGX via EODHD"]
    D --> E["Evaluate candidates"]
    E --> F{"New candidates?"}
    F -->|Yes| G["Create CANDIDATE entries"]
    F -->|No| H["Log: no new candidates"]
    G --> I["Update watchlist"]""",

    "DcaScheduler": """graph TD
    A["checkAndExecuteDca - 10:15 Daily"] --> B{"DCA enabled?"}
    B -->|No| C["Skip"]
    B -->|Yes| D["Check day of month"]
    D --> E{"Is execution day?"}
    E -->|No| F{"Weekend fallback?"}
    E -->|Yes| G["Execute NGX DCA"]
    F -->|Yes| G
    F -->|No| H["Skip - not DCA day"]
    G --> I["Execute US DCA"]""",

    "DividendCheckScheduler": """graph TD
    A["weeklyDividendCheck - MON 09:00"] --> B["Read alert config"]
    B --> C{"Track ex-dates?"}
    C -->|Yes| D["Check upcoming ex-dates"]
    C -->|No| E["Skip ex-date check"]
    D --> F{"Dates in alert window?"}
    F -->|Yes| G["Log upcoming ex-dates"]
    F -->|No| E
    E --> H{"Reinvestment enabled?"}
    G --> H
    H -->|Yes| I["Process reinvestments"]
    H -->|No| J["End"]""",

    "FundamentalUpdateScheduler": """graph TD
    A["weeklyFundamentalUpdate - SUN 20:00"] --> B["Get all core holdings"]
    B --> C{"Any holdings?"}
    C -->|No| D["Skip - none found"]
    C -->|Yes| E["Score each via FundamentalScreener"]
    E --> F["Log scores per holding"]""",

    "RebalanceScheduler": """graph TD
    A["quarterlyRebalanceCheck - Quarterly"] --> B["Refresh market values"]
    B --> C["Calculate current vs target weights"]
    C --> D{"Drift exceeds threshold?"}
    D -->|No| E["No rebalance needed"]
    D -->|Yes| F["Generate RebalanceActions"]
    F --> G{"Require approval?"}
    G -->|Yes| H["Save as PENDING"]
    G -->|No| I["Execute rebalance trades"]""",

    "CircuitBreaker": """graph TD
    A["periodicCircuitCheck - every 10m"] --> B["Check daily P&L"]
    B --> C{"Daily loss >= 5%?"}
    C -->|Yes| D["HALT satellite - daily"]
    C -->|No| E["Check weekly P&L"]
    E --> F{"Weekly loss >= 10%?"}
    F -->|Yes| G["HALT until Monday"]
    F -->|No| H["Check cross-market"]
    H --> I{"Cross loss >= 3%?"}
    I -->|Yes| J["HALT satellite"]
    I -->|No| K["All clear"]""",

    "SettlementCashTracker": """graph TD
    A["dailySettlementProcessing - 09:00"] --> B["Get today's date"]
    B --> C["Check settling entries per market"]
    C --> D{"Settlement date reached?"}
    D -->|Yes| E["Move to available cash pool"]
    D -->|No| F["No settlements today"]
    E --> G["Log processed settlements"]""",

    "StopLossMonitor": """graph TD
    A["monitorNgxStops - every 5m"] --> B["Get all open positions"]
    B --> C["Filter NGX by .XNSA suffix"]
    C --> D{"For each position"}
    D --> E["Get latest close price"]
    E --> F{"Price <= stop-loss?"}
    F -->|Yes| G["Mark for closure"]
    F -->|No| H["Update price & P&L"]
    G --> D
    H --> D""",

    "PortfolioReconciler": """graph TD
    A["dailyReconciliation - 09:30"] --> B["Get DB open positions"]
    B --> C["Get broker holdings via Gateway"]
    C --> D["Compare DB vs Broker"]
    D --> E{"Mismatches?"}
    E -->|No| F["Reconciliation passed"]
    E -->|Yes| G["Activate Kill Switch"]
    G --> H["Send urgent notification"]""",

    "CashReconciler": """graph TD
    A["dailyCashReconciliation - 09:35"] --> B["Reconcile NGX cash"]
    B --> C["Get internal cash from tracker"]
    C --> D["Get broker cash via Gateway"]
    D --> E{"Mismatch > 2%?"}
    E -->|Yes| F["Activate Kill Switch"]
    E -->|No| G["NGX OK"]
    F --> H["Send urgent notification"]
    G --> I["Reconcile US cash"]
    H --> I""",
}

# ═══════════════════════════════════════════════════════════════════════════════
# CROSS-CUTTING MERMAID DIAGRAMS
# ═══════════════════════════════════════════════════════════════════════════════

DIAGRAM_TIMELINE_MORNING = """gantt
    title Weekday Morning Schedule 06:00-14:30 WAT
    dateFormat YYYY-MM-DD HH:mm
    axisFormat %H:%M
    section News
    Early Morning News       :2024-01-01 08:00, 20m
    section Reconciliation
    Settlement Processing    :2024-01-01 09:00, 15m
    Pre-market RSS           :2024-01-01 09:00, 15m
    Portfolio Reconcile      :2024-01-01 09:30, 5m
    Cash Reconcile           :2024-01-01 09:35, 5m
    section Long-term
    DCA Check                :2024-01-01 10:15, 10m
    section Risk
    NGX Stop-Loss 5m loop    :2024-01-01 10:00, 270m
    Circuit Breaker 10m loop :2024-01-01 10:00, 270m
    section News 2
    Mid-day News             :2024-01-01 12:30, 20m"""

DIAGRAM_TIMELINE_AFTERNOON = """gantt
    title Weekday Afternoon-Evening Schedule 14:30-22:00 WAT
    dateFormat YYYY-MM-DD HH:mm
    axisFormat %H:%M
    section Data
    EODHD Data Pull          :2024-01-01 15:00, 15m
    NGX Backup Scrape        :2024-01-01 15:05, 10m
    ETF NAV Scrape           :2024-01-01 16:00, 15m
    section Signals
    NGX Signal Generation    :2024-01-01 15:15, 15m
    section Risk
    US Stop-Loss 5m loop     :2024-01-01 15:00, 360m
    Circuit Breaker 10m loop :2024-01-01 14:30, 450m
    section News
    After-market News        :2024-01-01 17:00, 20m
    section Discovery
    Daily Observation        :2024-01-01 20:00, 15m
    Daily Demotion           :2024-01-01 20:30, 15m
    section AI and Signals
    Cross-article Synthesis  :2024-01-01 21:00, 15m
    US Signal Generation     :2024-01-01 21:30, 15m"""

DIAGRAM_TIMELINE_WEEKEND = """gantt
    title Weekend Schedule WAT
    dateFormat YYYY-MM-DD HH:mm
    axisFormat %H:%M
    section Sunday
    Weekly Screener Scan     :2024-01-01 06:00, 30m
    Fundamental Update       :2024-01-01 20:00, 60m
    section Monday
    Dividend Check           :2024-01-01 09:00, 15m"""

DIAGRAM_DATA_FLOW = """graph LR
    subgraph Collection
        DC["DataCollection"]
        NS["NewsScraper"]
    end
    subgraph Analysis
        AI["AiAnalysis"]
        DISC["Discovery"]
    end
    subgraph Signals
        SIG["SignalGeneration"]
    end
    subgraph Execution
        DCA["DCA"]
        DIV["Dividend"]
        REB["Rebalance"]
    end
    subgraph Risk
        CB["CircuitBreaker"]
        SL["StopLoss"]
        SCT["Settlement"]
        PR["Reconciler"]
        CR["CashReconciler"]
    end
    DC -->|"OHLCV bars"| SIG
    DC -->|"OHLCV bars"| SL
    NS -->|"News articles"| AI
    AI -->|"Sentiment"| SIG
    DISC -->|"Watchlist"| SIG
    SIG -->|"Trade signals"| DCA
    SCT -->|"Cash ledger"| CR
    PR -->|"Kill switch"| SIG
    CR -->|"Kill switch"| SIG"""

DIAGRAM_DEPENDENCY = """graph TD
    KS["KillSwitchService"]
    NR["NotificationRouter"]
    BG["BrokerGateway"]
    SCT["SettlementCashTracker"]
    PR["PortfolioReconciler"] --> KS
    PR --> NR
    PR --> BG
    CR["CashReconciler"] --> KS
    CR --> NR
    CR --> BG
    CR --> SCT"""

DIAGRAM_KILL_SWITCH = """sequenceDiagram
    participant PR as PortfolioReconciler
    participant BG as BrokerGateway
    participant KS as KillSwitchService
    participant NR as NotificationRouter
    PR->>BG: Get broker holdings
    BG-->>PR: Holdings list
    PR->>PR: Compare DB vs Broker
    alt Mismatch found
        PR->>KS: activate(reason)
        PR->>NR: sendUrgent(alert)
        Note over KS: All trading HALTED
    else Match
        PR->>PR: Log success
    end"""

DIAGRAM_SETTLEMENT = """graph LR
    A["Sale Executed"] --> B["Record in Settling Pool"]
    B --> C{"Settlement Date?"}
    C -->|"Not yet T+2/T+1"| D["Remain in Settling"]
    C -->|"Date reached"| E["Move to Available Cash"]
    E --> F["Available for New Trades"]
    D --> C"""

DIAGRAM_SIGNAL_EXEC = """sequenceDiagram
    participant DC as DataCollection
    participant AI as AiAnalysis
    participant SIG as SignalGeneration
    participant KS as KillSwitch
    participant CB as CircuitBreaker
    participant EX as OrderRouter
    DC->>DC: Collect OHLCV and News
    AI->>AI: Analyze sentiment
    SIG->>SIG: Generate signals
    SIG->>KS: Check kill switch
    alt Kill switch active
        KS-->>SIG: BLOCKED
    else OK
        SIG->>CB: Check circuit breaker
        alt Circuit broken
            CB-->>SIG: HALTED
        else OK
            SIG->>EX: Route order
        end
    end"""

DIAGRAM_DISCOVERY_LIFECYCLE = """stateDiagram-v2
    [*] --> CANDIDATE: Weekly screener
    CANDIDATE --> OBSERVATION: 2+ signals
    OBSERVATION --> PROMOTED: Meets criteria
    OBSERVATION --> DEMOTED: No signals N days
    PROMOTED --> DEMOTED: Low fundamentals
    DEMOTED --> CANDIDATE: Cooldown expires
    PROMOTED --> [*]: Removed"""

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION GENERATORS
# ═══════════════════════════════════════════════════════════════════════════════

def _all_jobs():
    """Flatten all jobs with scheduler context attached."""
    idx = 0
    for s in SCHEDULERS:
        for j in s["jobs"]:
            idx += 1
            yield idx, s, j


def _total_jobs():
    return sum(len(s["jobs"]) for s in SCHEDULERS)


def build_title_and_toc(doc):
    """Sections: Title page + Table of Contents."""
    add_title_page(doc, "Cron Jobs Reference",
                   "Scheduled Task Architecture & Documentation",
                   "Engineering Team, DevOps, QA")
    doc.add_heading("Table of Contents", level=1)
    add_toc_field(doc)
    doc.add_page_break()


def build_executive_summary(doc):
    """Section 2: Executive Summary."""
    doc.add_heading("Executive Summary", level=1)
    total = _total_jobs()
    doc.add_paragraph(
        f"The NGX Trading Bot operates {total} scheduled tasks across "
        f"{len(SCHEDULERS)} Java classes organized into 6 functional modules. "
        f"All cron-based tasks use the Africa/Lagos timezone (WAT, UTC+1)."
    )
    doc.add_paragraph("")
    doc.add_heading("Scheduling Patterns", level=2)
    add_styled_table(doc,
        ["Pattern", "Count", "Details"],
        [
            ["Cron-based", str(total - 1), "Time/day-specific triggers"],
            ["fixedDelay", "1", "AiAnalysisScheduler.processUnanalyzedArticles (15 min)"],
        ])
    doc.add_heading("Category Breakdown", level=2)
    cats = {}
    for s in SCHEDULERS:
        c = s["category"]
        cats[c] = cats.get(c, 0) + len(s["jobs"])
    add_styled_table(doc,
        ["Category", "Schedulers", "Jobs"],
        [[cat, str(sum(1 for s in SCHEDULERS if s["category"] == cat)), str(cnt)]
         for cat, cnt in cats.items()])
    doc.add_paragraph("")
    add_callout(doc, "All scheduled times are in WAT (West Africa Time, UTC+1). "
                "The bot does NOT trade outside 10:00-14:30 WAT for NGX.", BRAND_ORANGE)
    doc.add_page_break()


def build_master_inventory(doc):
    """Section 3: Master Job Inventory Table (landscape)."""
    switch_to_landscape(doc)
    doc.add_heading("Master Job Inventory", level=1)
    doc.add_paragraph("Complete inventory of all scheduled tasks sorted by execution time.")
    rows = []
    for idx, s, j in _all_jobs():
        rows.append([
            str(idx), s["name"], j["method"], j["cron"],
            j["time"], j["days"], s["category"], j["purpose"]
        ])
    add_styled_table(doc,
        ["#", "Scheduler", "Method", "Schedule", "Time (WAT)", "Days", "Category", "Purpose"],
        rows,
        col_widths=[0.3, 1.5, 1.5, 1.3, 1.0, 0.7, 0.8, 2.4])
    switch_to_portrait(doc)
    doc.add_page_break()


def build_daily_timeline(doc):
    """Section 4: Daily Timeline with gantt charts."""
    doc.add_heading("Daily Execution Timeline", level=1)
    doc.add_paragraph("Visual timeline of all scheduled tasks across a typical trading day.")
    doc.add_heading("Weekday Morning (06:00-14:30 WAT)", level=2)
    add_diagram(doc, DIAGRAM_TIMELINE_MORNING, "Morning schedule: reconciliation, news, risk monitoring")
    doc.add_heading("Weekday Afternoon-Evening (14:30-22:00 WAT)", level=2)
    add_diagram(doc, DIAGRAM_TIMELINE_AFTERNOON, "Afternoon/evening: data pull, signals, discovery, AI")
    doc.add_heading("Weekend & Weekly Jobs", level=2)
    add_diagram(doc, DIAGRAM_TIMELINE_WEEKEND, "Weekend: Sunday screener + fundamentals, Monday dividends")
    doc.add_paragraph("")
    add_callout(doc, "Quarterly job (RebalanceScheduler) runs on 1st of Jan/Apr/Jul/Oct at 15:00 WAT.",
                BRAND_GREEN)
    doc.add_page_break()


def build_scheduler_sections(doc):
    """Sections 5-18: Per-Scheduler detailed documentation."""
    for i, s in enumerate(SCHEDULERS):
        sec_num = i + 5
        doc.add_heading(f"{sec_num}. {s['display']}", level=1)
        doc.add_paragraph(s["overview"])
        add_callout(doc, f"File: {s['path']}", BRAND_GRAY)

        # Dependencies table
        if s["deps"]:
            doc.add_heading("Dependencies", level=2)
            add_styled_table(doc,
                ["Type", "Name", "Purpose"],
                [[d[0], d[1], d[2]] for d in s["deps"]])

        # DB tables
        if s["tables_r"] or s["tables_w"]:
            doc.add_heading("Database Tables", level=2)
            db_rows = []
            for t in s["tables_r"]:
                db_rows.append([t, "READ", ""])
            for t in s["tables_w"]:
                mode = "READ/WRITE" if t in s["tables_r"] else "WRITE"
                if t not in s["tables_r"]:
                    db_rows.append([t, mode, ""])
            add_styled_table(doc, ["Table", "Access", "Notes"], db_rows)

        # Kill switch
        if s["kill_switch"]:
            add_callout(doc, "WARNING: This scheduler can ACTIVATE the kill switch, "
                        "halting ALL trading.", BRAND_RED)

        # Per-job subsections
        for j in s["jobs"]:
            doc.add_heading(f"{j['method']}()", level=3)
            add_callout(doc, f"Schedule: {j['cron']}  |  Time: {j['time']}  |  Days: {j['days']}")
            doc.add_paragraph(j["purpose"])

        # Flowchart
        if s["name"] in FLOWCHARTS:
            doc.add_heading("Execution Flow", level=2)
            add_diagram(doc, FLOWCHARTS[s["name"]],
                        f"{s['display']} - primary job flow", width=5.5)

        add_section_divider(doc)
        doc.add_page_break()


def build_data_flow(doc):
    """Section 19: Inter-Module Data Flow."""
    doc.add_heading("19. Inter-Module Data Flow", level=1)
    doc.add_paragraph(
        "Data flows through a pipeline: Collection gathers raw data, Analysis extracts "
        "insights, Signal Generation produces trade decisions, and Risk Management gates "
        "execution."
    )
    add_diagram(doc, DIAGRAM_DATA_FLOW, "Cross-module data flow pipeline", width=6.5)
    doc.add_heading("Data Flow Details", level=2)
    add_styled_table(doc,
        ["Source Job", "Data Object", "Consumer Job"],
        [
            ["DataCollection.pullEodhdData", "ohlcv_bar records", "SignalGeneration, StopLossMonitor"],
            ["DataCollection.preMarketNewsScan", "news_item records", "AiAnalysis"],
            ["NewsScraper (3 scans)", "news_item records", "AiAnalysis"],
            ["AiAnalysis.processUnanalyzed", "ai_analysis (sentiment)", "SignalGeneration"],
            ["AiAnalysis.crossArticle", "synthesis summaries", "SignalGeneration"],
            ["Discovery.weeklyScreener", "discovered_stock (CANDIDATE)", "Discovery.observation"],
            ["Discovery.observation", "discovered_stock (PROMOTED)", "SignalGeneration"],
            ["SignalGeneration (NGX/US)", "trade_signal entities", "Execution layer"],
            ["SettlementCashTracker", "CashLedger (available)", "CashReconciler"],
            ["PortfolioReconciler", "Kill switch activation", "All execution"],
        ])
    doc.add_page_break()


def build_dependency_graph(doc):
    """Section 20: Dependency Graph."""
    doc.add_heading("20. Shared Service Dependency Graph", level=1)
    doc.add_paragraph(
        "Several critical services are shared across schedulers. KillSwitchService "
        "and NotificationRouter are the most cross-cutting dependencies."
    )
    add_diagram(doc, DIAGRAM_DEPENDENCY, "Shared service dependency graph", width=5.0)
    doc.add_heading("Shared Service Usage", level=2)
    add_styled_table(doc,
        ["Shared Service", "Used By"],
        [
            ["KillSwitchService", "PortfolioReconciler, CashReconciler"],
            ["NotificationRouter", "PortfolioReconciler, CashReconciler"],
            ["BrokerGateway", "PortfolioReconciler, CashReconciler"],
            ["SettlementCashTracker", "CashReconciler (reads ledger), Execution layer (records trades)"],
            ["PositionRepository", "PortfolioReconciler, StopLossMonitor"],
            ["OhlcvRepository", "StopLossMonitor, SignalGeneration (indirect)"],
            ["CompositeSignalScorer", "SignalGenerationScheduler"],
            ["AiCostTracker", "AiAnalysisScheduler (budget gating)"],
        ])
    doc.add_page_break()


def build_risk_architecture(doc):
    """Section 21: Risk & Safety Architecture."""
    doc.add_heading("21. Risk & Safety Architecture", level=1)
    doc.add_paragraph(
        "The trading bot implements multiple safety layers: kill switch for hard stops, "
        "circuit breakers for loss limits, and settlement tracking for cash management."
    )

    doc.add_heading("Kill Switch Interaction", level=2)
    doc.add_paragraph(
        "Only PortfolioReconciler and CashReconciler can activate the kill switch. "
        "Once active, ALL trading halts until manual deactivation."
    )
    add_diagram(doc, DIAGRAM_KILL_SWITCH, "Kill switch activation sequence", width=5.5)

    doc.add_heading("Settlement Cash Flow", level=2)
    doc.add_paragraph(
        "Sales enter a settling pool. NGX settles T+2, US settles T+1. "
        "SettlementCashTracker processes at 09:00 daily."
    )
    add_diagram(doc, DIAGRAM_SETTLEMENT, "Settlement flow: sale to available cash", width=5.5)

    doc.add_heading("Signal-to-Execution Pipeline", level=2)
    doc.add_paragraph(
        "Trade signals must pass through kill switch and circuit breaker checks "
        "before reaching the order router."
    )
    add_diagram(doc, DIAGRAM_SIGNAL_EXEC, "Signal to execution sequence", width=5.5)

    doc.add_heading("Discovery Lifecycle", level=2)
    doc.add_paragraph(
        "Discovered stocks progress through states: CANDIDATE (new), OBSERVATION "
        "(promising), PROMOTED (active), DEMOTED (cooldown)."
    )
    add_diagram(doc, DIAGRAM_DISCOVERY_LIFECYCLE, "Stock discovery state machine", width=5.0)
    doc.add_page_break()


def build_reference_card(doc):
    """Section 22: Schedule Reference Card."""
    doc.add_heading("22. Schedule Reference Card", level=1)
    doc.add_paragraph("Compact reference of all scheduled tasks sorted by execution time.")

    # Sort jobs by time for the reference card
    all_jobs = list(_all_jobs())
    def sort_key(item):
        _, s, j = item
        t = j["time"]
        if t.startswith("Every"):
            return "99:" + t
        # Extract HH:MM for sorting
        parts = t.split(":")
        if len(parts) == 2 and parts[0].isdigit():
            return t
        return "50:" + t
    all_jobs.sort(key=sort_key)

    rows = []
    for idx, s, j in all_jobs:
        rows.append([
            j["time"], j["days"], s["name"].replace("Scheduler", ""),
            j["method"], s["category"]
        ])
    add_styled_table(doc,
        ["Time (WAT)", "Days", "Scheduler", "Method", "Category"],
        rows,
        col_widths=[1.2, 0.8, 1.8, 2.0, 1.0])

    doc.add_paragraph("")
    add_callout(doc, f"Total: {_total_jobs()} scheduled tasks across "
                f"{len(SCHEDULERS)} classes in 6 modules.", BRAND_GREEN)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("  Generating CRON_JOBS_REFERENCE.docx")
    print("=" * 60)

    doc = Document()
    setup_styles(doc)
    apply_header_footer(doc.sections[0])

    print("  [1/10] Title page & TOC...")
    build_title_and_toc(doc)

    print("  [2/10] Executive summary...")
    build_executive_summary(doc)

    print("  [3/10] Master inventory table (landscape)...")
    build_master_inventory(doc)

    print("  [4/10] Daily timeline diagrams...")
    build_daily_timeline(doc)

    print("  [5/10] Per-scheduler sections (14 schedulers)...")
    build_scheduler_sections(doc)

    print("  [6/10] Inter-module data flow...")
    build_data_flow(doc)

    print("  [7/10] Dependency graph...")
    build_dependency_graph(doc)

    print("  [8/10] Risk & safety architecture...")
    build_risk_architecture(doc)

    print("  [9/10] Reference card...")
    build_reference_card(doc)

    print("  [10/10] Saving document...")
    save_doc(doc, "CRON_JOBS_REFERENCE.docx")

    print("=" * 60)
    print(f"  Done! {_total_jobs()} jobs across {len(SCHEDULERS)} schedulers documented.")
    print("=" * 60)


if __name__ == "__main__":
    main()
