"""
NGX Trading Bot — Word Document Generator
Generates 7 professional .docx files with styled content, tables, and Mermaid diagrams.

Usage:
    python generate_docs.py

Output:
    docs/word/DEVELOPER_GUIDE.docx
    docs/word/QA_GUIDE.docx
    docs/word/BUSINESS_OVERVIEW.docx
    docs/word/PRODUCT_SPEC.docx
    docs/word/PITCH.docx
    docs/word/API_REFERENCE.docx
    docs/word/DEPLOYMENT_GUIDE.docx
"""

import os
import base64
import io
import time
from pathlib import Path

import requests
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Configuration ───────────────────────────────────────────────────────────
OUTPUT_DIR = Path(__file__).parent / "word"
BRAND_BLUE = RGBColor(0x33, 0x9A, 0xF0)        # #339AF0
BRAND_GREEN = RGBColor(0x51, 0xCF, 0x66)        # #51CF66
BRAND_RED = RGBColor(0xF0, 0x3E, 0x3E)          # #F03E3E
BRAND_ORANGE = RGBColor(0xFF, 0xA9, 0x4D)       # #FFA94D
BRAND_DARK = RGBColor(0x21, 0x25, 0x29)         # #212529
BRAND_GRAY = RGBColor(0x86, 0x8E, 0x96)         # #868E96
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = RGBColor(0xE7, 0xF5, 0xFF)      # #E7F5FF


# ─── Mermaid Diagram Renderer ────────────────────────────────────────────────

def render_mermaid(mermaid_code: str, max_retries: int = 2) -> bytes | None:
    """Render a Mermaid diagram to PNG via mermaid.ink API."""
    encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("utf-8")
    url = f"https://mermaid.ink/img/{encoded}?type=png&bgColor=white"

    for attempt in range(max_retries + 1):
        try:
            resp = requests.get(url, timeout=30)
            if resp.status_code == 200 and len(resp.content) > 100:
                return resp.content
        except requests.RequestException:
            pass
        if attempt < max_retries:
            time.sleep(2)

    print(f"  [WARNING] Failed to render Mermaid diagram after {max_retries + 1} attempts")
    return None


# ─── Document Styling Helpers ─────────────────────────────────────────────────

def setup_styles(doc: Document):
    """Configure document-wide styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = BRAND_DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, bold, color) in {
        "Heading 1": (Pt(24), True, BRAND_BLUE),
        "Heading 2": (Pt(18), True, BRAND_DARK),
        "Heading 3": (Pt(14), True, BRAND_BLUE),
        "Heading 4": (Pt(12), True, BRAND_DARK),
    }.items():
        h = doc.styles[level]
        h.font.name = "Calibri"
        h.font.size = size
        h.font.bold = bold
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(18) if level != "Heading 1" else Pt(6)
        h.paragraph_format.space_after = Pt(8)


def add_title_page(doc: Document, title: str, subtitle: str, audience: str):
    """Add a styled title page."""
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = BRAND_BLUE
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_GRAY
    run.font.name = "Calibri"

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Audience: {audience}")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = BRAND_DARK

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NGX Trading Bot")
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_BLUE
    run.font.bold = True

    doc.add_page_break()


def add_styled_table(doc: Document, headers: list[str], rows: list[list[str]],
                     col_widths: list[float] | None = None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="339AF0"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            # Alternate row shading
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F3F5"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    doc.add_paragraph("")
    return table


def add_diagram(doc: Document, mermaid_code: str, caption: str = "",
                width: float = 6.0):
    """Render a Mermaid diagram and embed it in the document."""
    img_data = render_mermaid(mermaid_code)
    if img_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(img_data), width=Inches(width))

        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = BRAND_GRAY
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Diagram: {caption}]")
        run.font.italic = True
        run.font.color.rgb = BRAND_RED


def add_code_block(doc: Document, code: str, language: str = ""):
    """Add a styled code block."""
    p = doc.add_paragraph()
    if language:
        run = p.add_run(f"  {language}")
        run.font.size = Pt(8)
        run.font.color.rgb = BRAND_GRAY
        run.font.italic = True
        p = doc.add_paragraph()

    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_DARK
    # Light background
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
    p._p.get_or_add_pPr().append(shading)
    p.paragraph_format.left_indent = Cm(0.5)


def add_callout(doc: Document, text: str, color: RGBColor = BRAND_BLUE):
    """Add a styled callout/note box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("  " + text)
    run.font.size = Pt(10)
    run.font.color.rgb = color
    run.font.italic = True
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
    p._p.get_or_add_pPr().append(shading)


def add_bullet(doc: Document, text: str, bold_prefix: str = ""):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        p.runs[0].font.size = Pt(11) if p.runs else None
        if not p.runs:
            run = p.add_run(text)
            run.font.size = Pt(11)
        else:
            p.text = text


def add_section_divider(doc: Document):
    """Add a visual divider between sections."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xDE, 0xE2, 0xE6)
    run.font.size = Pt(8)


def save_doc(doc: Document, filename: str):
    """Save document to the output directory."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc.save(str(path))
    print(f"  Saved: {path}")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1: DEVELOPER GUIDE
# ═══════════════════════════════════════════════════════════════════════════════

def generate_developer_guide():
    print("\n[1/7] Generating DEVELOPER_GUIDE.docx...")
    doc = Document()
    setup_styles(doc)

    add_title_page(doc, "Developer Guide", "Technical Architecture & Setup", "Software engineers onboarding or contributing")

    # ── Architecture Overview ──
    doc.add_heading("Architecture Overview", level=1)
    doc.add_paragraph(
        "The NGX Trading Bot is a Spring Boot 3.3 application (Java 17) that autonomously trades stocks on the "
        "Nigerian Stock Exchange (NGX) and US equity markets through the Trove/Meritrade brokerage platform. "
        "It uses Playwright-Java for browser automation, WAHA for WhatsApp notifications, and PostgreSQL 16 for persistence."
    )

    add_styled_table(doc, ["Component", "Technology", "Purpose"], [
        ["Runtime", "Java 17, Spring Boot 3.3.6", "Application framework"],
        ["Database", "PostgreSQL 16 + Flyway", "Persistence + schema migrations"],
        ["Browser Automation", "Playwright-Java 1.49.0", "Meritrade broker interaction"],
        ["WhatsApp", "WAHA (Docker)", "Primary notification channel"],
        ["Telegram", "Bot API", "Fallback notification channel"],
        ["Market Data", "EODHD API", "OHLCV bars + fundamentals"],
        ["AI Analysis", "Claude API (Anthropic)", "News sentiment + earnings analysis"],
        ["Web Scraping", "Jsoup", "News from 6 sources"],
        ["Build", "Maven", "Dependency management + build"],
    ])

    add_diagram(doc, """graph TB
    subgraph DATA["Data Layer"]
        EODHD["EODHD API\\nMarket Data"]
        NEWS["News Scrapers\\n6 Sources"]
        BROKER["Trove Browser\\nPlaywright"]
    end

    subgraph ENGINE["Signal Engine"]
        IND["Technical Indicators\\nRSI, MACD, ATR, SMA"]
        STRAT["10 Strategies\\nCORE + SATELLITE"]
        AI["Claude AI\\nSentiment Analysis"]
    end

    subgraph RISK["Risk & Execution"]
        RM["Risk Manager\\n7 Hard Rules"]
        CB["Circuit Breaker\\nDaily + Weekly"]
        KS["Kill Switch"]
        OR["Order Router"]
    end

    subgraph EXEC["Execution"]
        BA["Browser Agent\\nPlaywright"]
        OTP["OTP Handler\\nGmail + WhatsApp"]
        SL["Session Lock"]
    end

    subgraph NOTIFY["Notifications"]
        WA["WhatsApp\\nWAHA"]
        TG["Telegram\\nBot API"]
        APR["Trade Approval\\n5-min timeout"]
    end

    DATA --> ENGINE --> RISK --> EXEC --> NOTIFY
    """, "Figure 1: System Architecture", 5.5)

    # ── Module Breakdown ──
    doc.add_heading("Module Breakdown", level=1)

    modules = [
        ("common", "Shared models (TradeSide, SignalStrength, TradeSignal), exceptions (KillSwitchActiveException), base entities"),
        ("data", "EODHD API client, OHLCV bar entity/repository, market data fetching"),
        ("signal", "RSI, MACD, ATR, SMA/EMA calculators, CompositeSignalScorer, TradeSignal generation"),
        ("strategy", "10 trading strategies, Strategy interface, StrategyPool (CORE/SATELLITE), StrategyMarket (NGX/US/BOTH)"),
        ("risk", "RiskManager (5 checks), CircuitBreaker (daily/weekly), SettlementCashTracker (T+2 NGX, T+1 US), Position/PortfolioSnapshot entities"),
        ("execution", "OrderRouter, TroveBrowserAgent (Playwright), BrokerGateway, OtpHandler + EmailOtpReader, KillSwitchService, BrowserSessionLock, OrderRecoveryService"),
        ("notification", "NotificationRouter (WhatsApp + Telegram), MessageFormatter, TradeApprovalService, WhatsAppWebhookController"),
        ("news", "6 scrapers (BusinessDay, Nairametrics, SeekingAlpha, Reuters, CBN, NGX), NewsEventClassifier, EventImpactRules"),
        ("ai", "Claude API client, AiCostTracker, budget enforcement ($5/day, $100/month), deep analysis triggers"),
        ("discovery", "Stock screener (EODHD), watchlist management, observation pipeline, promotion/demotion policies"),
        ("longterm", "DCA scheduling, dividend tracking + reinvestment, portfolio rebalancing (quarterly, 10% drift)"),
        ("backtest", "BacktestRunner, SimulatedOrderExecutor, PerformanceAnalyzer (Sharpe, drawdown, win rate)"),
        ("dashboard", "REST API controllers (29 endpoints), PortfolioReconciler, CashReconciler, FxRateReconciler"),
        ("config", "PlaywrightConfig, *Properties classes, YAML-driven configuration"),
        ("scheduling", "Cron jobs for market hours, settlement processing, screenshot cleanup"),
    ]

    add_styled_table(doc, ["Module", "Responsibility"], [[m, r] for m, r in modules], [1.5, 5.0])

    # ── Deployment Architecture ──
    doc.add_heading("Deployment Architecture", level=1)

    add_diagram(doc, """graph LR
    subgraph DOCKER["Docker Compose Stack"]
        BOT["trading-bot\\n:8080\\nSpring Boot + Playwright"]
        PG["postgres\\n:5432\\nPostgreSQL 16"]
        WAHA["waha\\n:3000\\nWhatsApp WEBJS"]
    end

    subgraph EXT["External Services"]
        EODHD["EODHD API"]
        TROVE["Trove Broker"]
        TGAPI["Telegram API"]
        CLAUDE["Anthropic API"]
        GMAIL["Gmail IMAP"]
    end

    BOT --> PG
    BOT --> WAHA
    BOT --> EODHD
    BOT --> TROVE
    BOT --> TGAPI
    BOT --> CLAUDE
    BOT --> GMAIL
    """, "Figure 2: Deployment Architecture", 5.5)

    # ── Trade Execution Sequence ──
    doc.add_heading("Trade Execution Sequence", level=1)

    add_diagram(doc, """sequenceDiagram
    participant SCH as Scheduler
    participant SIG as SignalEngine
    participant STR as Strategy
    participant RM as RiskManager
    participant KS as KillSwitch
    participant APR as ApprovalService
    participant WA as WhatsApp
    participant OR as OrderRouter
    participant BA as BrowserAgent
    participant TROVE as Trove Broker

    SCH->>SIG: triggerSignalGeneration()
    SIG->>SIG: calculateIndicators(RSI,MACD,ATR)
    SIG->>STR: evaluate(indicators, bars)
    STR-->>SIG: TradeSignal(BUY, ZENITHBANK, 48.00)
    SIG->>KS: checkOrThrow()
    KS-->>SIG: OK (not active)
    SIG->>RM: validateAll(signal, portfolio)
    RM-->>SIG: PASS (all 5 checks)
    SIG->>APR: requestApproval(signal)
    APR->>WA: Send approval request
    WA-->>APR: User replies YES
    APR->>OR: route(approvedOrder)
    OR->>BA: submitOrder(LIMIT, BUY, ZENITHBANK, 48.00, 100)
    BA->>TROVE: Browser automation (fill form, click confirm)
    TROVE-->>BA: Order confirmed
    BA-->>OR: orderId=TRD-12345
    OR->>WA: Send confirmation notification
    """, "Figure 3: Trade Execution Sequence", 6.0)

    # ── Database ──
    doc.add_heading("Database", level=1)
    doc.add_paragraph("30 Flyway migrations in src/main/resources/db/migration/ (V1 through V30):")

    migrations = [
        ["V1", "ohlcv_bars"], ["V2", "etf_valuations"], ["V3", "trade_orders"],
        ["V4", "positions"], ["V5", "portfolio_snapshots"], ["V6", "corporate_actions"],
        ["V7", "market_indices"], ["V8", "watchlist_stocks"], ["V9", "trade_signals"],
        ["V10", "notification_log"], ["V11", "kill_switch_state"], ["V12", "approval_requests"],
        ["V13", "news_items"], ["V14", "broker_sessions"], ["V15", "risk_events"],
        ["V16", "circuit_breaker_log"], ["V17", "sector_mapping"], ["V18", "system_config"],
        ["V19", "updated_at triggers"], ["V20", "Seed watchlist stocks"],
        ["V21", "ai_analysis, ai_cost_ledger"], ["V22", "News event columns"],
        ["V23", "discovered_stocks"], ["V24", "discovery_events"], ["V25", "core_holdings"],
        ["V26", "dca_plans"], ["V27", "dividend_events"], ["V28", "rebalance_actions"],
        ["V29", "Pool/market/currency on positions"], ["V30", "backtest_runs, backtest_trades, equity_curve_points"],
    ]
    add_styled_table(doc, ["Migration", "Table(s) Created"], migrations, [1.0, 5.5])

    doc.add_heading("Entity Relationship Diagram", level=2)
    add_diagram(doc, """erDiagram
    watchlist_stocks ||--o{ ohlcv_bars : "market data for"
    ohlcv_bars ||--o{ trade_signals : "generates"
    trade_signals ||--o{ trade_orders : "triggers"
    trade_orders ||--o{ positions : "creates"
    positions }o--|| portfolio_snapshots : "aggregates to"
    news_items ||--o{ ai_analysis : "analyzed by"
    discovered_stocks ||--o{ discovery_events : "audit log"
    discovered_stocks ||--o{ watchlist_stocks : "promotes to"
    dca_plans ||--o{ trade_orders : "executes"
    dividend_events ||--o{ trade_orders : "reinvests via"
    backtest_runs ||--o{ backtest_trades : "contains"
    backtest_runs ||--o{ equity_curve_points : "tracks"

    watchlist_stocks {
        bigint id PK
        varchar symbol
        varchar status
        varchar market
    }
    trade_orders {
        bigint id PK
        varchar order_id
        varchar symbol
        varchar side
        integer quantity
        numeric intended_price
        varchar status
    }
    positions {
        bigint id PK
        varchar symbol
        integer quantity
        numeric avg_entry_price
        varchar pool
        varchar market
        boolean is_open
    }
    """, "Figure 4: Entity Relationship Diagram", 5.5)

    # ── Local Setup ──
    doc.add_heading("Local Development Setup", level=1)

    doc.add_heading("Prerequisites", level=2)
    add_styled_table(doc, ["Requirement", "Version"], [
        ["Java", "17+ (Temurin recommended)"],
        ["Maven", "3.9+"],
        ["PostgreSQL", "16 (via Docker)"],
        ["Docker + Docker Compose", "v2.20+"],
        ["Node.js", "Not required (no frontend)"],
    ])

    doc.add_heading("Setup Steps", level=2)
    doc.add_paragraph("1. Clone the repository and configure environment variables:")
    add_code_block(doc, "cp .env.example .env\n# Edit .env with your API keys and credentials", "bash")
    doc.add_paragraph("2. Start infrastructure services:")
    add_code_block(doc, "docker compose up -d postgres waha", "bash")
    doc.add_paragraph("3. Scan WhatsApp QR code at http://localhost:3000")
    doc.add_paragraph("4. Build and run:")
    add_code_block(doc, "mvn clean compile\nmvn spring-boot:run", "bash")

    # ── Configuration ──
    doc.add_heading("Configuration Guide", level=1)
    doc.add_paragraph("All configuration lives in application.yml with profile overrides:")

    add_styled_table(doc, ["Profile", "File", "Use"], [
        ["(default)", "application.yml", "Local development"],
        ["prod", "application-prod.yml", "Docker production deployment"],
        ["integration", "application-integration.yml", "Integration test environment"],
    ])

    doc.add_heading("Key Configuration Sections", level=2)
    add_styled_table(doc, ["Section", "Controls"], [
        ["trading", "Market hours, timezone, watchlist symbols"],
        ["risk", "All 7 risk thresholds (percentages stored as decimals)"],
        ["meritrade", "Broker URL, credentials, CSS selectors for Playwright"],
        ["otp", "Gmail IMAP settings for OTP reading"],
        ["notification", "WhatsApp (WAHA) + Telegram config"],
        ["ai", "Claude API, model selection, budget limits ($5/day, $100/month)"],
        ["discovery", "Screener thresholds, watchlist sizing"],
        ["longterm", "DCA budgets, dividend reinvestment, rebalance rules"],
        ["strategies", "Per-strategy enable/disable + parameters"],
    ])

    # ── Coding Conventions ──
    doc.add_heading("Coding Conventions", level=1)
    add_styled_table(doc, ["Rule", "Convention"], [
        ["Monetary values", "BigDecimal — never double or float"],
        ["Trade dates", "LocalDate"],
        ["Timestamps", 'ZonedDateTime with ZoneId.of("Africa/Lagos")'],
        ["Dependency injection", "Constructor injection only (no field @Autowired)"],
        ["Scheduled tasks", 'zone = "Africa/Lagos" on all @Scheduled'],
        ["Orders", "LIMIT only — never market orders"],
        ["Percentages", "Stored as decimals in config (2.0 = 2%)"],
        ["Boilerplate", "Lombok (@RequiredArgsConstructor, @Slf4j, @Builder, @Data)"],
    ])

    # ── Common Commands ──
    doc.add_heading("Common Commands", level=1)
    add_styled_table(doc, ["Command", "Purpose"], [
        ["mvn clean compile", "Build the project"],
        ["mvn test", "Run unit tests only"],
        ["mvn verify -Pintegration", "Run integration tests"],
        ["mvn spring-boot:run", "Run locally"],
        ["docker compose up -d", "Full stack (bot + postgres + waha)"],
        ["docker compose up -d postgres waha", "Infrastructure only"],
        ["docker compose logs -f trading-bot", "Follow bot logs"],
    ])

    save_doc(doc, "DEVELOPER_GUIDE.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2: QA GUIDE
# ═══════════════════════════════════════════════════════════════════════════════

def generate_qa_guide():
    print("\n[2/7] Generating QA_GUIDE.docx...")
    doc = Document()
    setup_styles(doc)
    add_title_page(doc, "QA Guide", "Test Architecture & Coverage", "QA engineers testing the NGX Trading Bot")

    # ── Test Architecture ──
    doc.add_heading("Test Architecture", level=1)
    add_styled_table(doc, ["Tier", "Framework", "Count", "Runner", "Tag"], [
        ["Unit Tests", "JUnit 5 + Mockito + AssertJ", "~181 methods", "Maven Surefire", "(none)"],
        ["Integration Tests", "JUnit 5 + Spring Boot Test", "11 step files", "Maven Failsafe", '@Tag("integration")'],
    ])

    doc.add_paragraph(
        "Unit tests use H2 in-memory database and mocked dependencies. "
        "Integration tests hit real PostgreSQL, real EODHD API, real WAHA, and real Trove browser."
    )

    doc.add_heading("Test Pyramid", level=2)
    add_diagram(doc, """graph TB
    subgraph pyramid[" "]
        direction TB
        E2E["E2E (Step 11)\\n1 test class, 10 steps\\nFull pipeline validation"]
        INT["Integration Tests (Steps 01-10)\\n10 test classes, ~60 tests\\nReal services, real data"]
        UNIT["Unit Tests\\n24 test classes, ~181 methods\\nMocked dependencies, H2 DB"]
    end

    style E2E fill:#ff6b6b,color:#fff
    style INT fill:#ffa94d,color:#fff
    style UNIT fill:#51cf66,color:#fff
    """, "Figure 1: Test Pyramid", 4.5)

    doc.add_heading("Integration Test Dependency Flow", level=2)
    add_diagram(doc, """graph LR
    S01["Step 01\\nPostgreSQL\\n+ Flyway"] --> S02["Step 02\\nEODHD\\nMarket Data"]
    S01 --> S03["Step 03\\nTelegram"]
    S01 --> S04["Step 04\\nWhatsApp\\n+ WAHA"]
    S02 --> S05["Step 05\\nNews\\nScrapers"]
    S02 --> S06["Step 06\\nClaude AI"]
    S04 --> S07["Step 07\\nMeritrade\\nLogin"]
    S07 --> S08["Step 08\\nPaper\\nTrade"]
    S02 --> S09["Step 09\\nBacktesting"]
    S01 --> S10["Step 10\\nDashboard\\nREST API"]
    S02 --> S11["Step 11\\nEnd-to-End"]
    S03 --> S11
    S04 --> S11
    S06 --> S11

    style S01 fill:#339af0,color:#fff
    style S07 fill:#f03e3e,color:#fff
    style S08 fill:#f03e3e,color:#fff
    style S11 fill:#ae3ec9,color:#fff
    """, "Figure 2: Integration Test Dependencies (Blue=infra, Red=browser, Purple=full pipeline)", 6.0)

    # ── Unit Test Coverage ──
    doc.add_heading("Unit Test Coverage", level=1)
    doc.add_paragraph("24 test classes covering all critical modules:")

    doc.add_heading("Risk & Safety", level=2)
    add_styled_table(doc, ["Class", "What It Tests"], [
        ["RiskManagerTest", "All 5 risk checks: max positions, single position size, sector exposure, cash reserve, risk per trade"],
        ["CircuitBreakerTest", "Daily 5% loss trigger, weekly 10% loss trigger, reset behavior"],
        ["SettlementCashTrackerTest", "T+2 NGX settlement, T+1 US settlement, overspend prevention, cross-market isolation"],
    ])

    doc.add_heading("Execution", level=2)
    add_styled_table(doc, ["Class", "What It Tests"], [
        ["OrderRouterTest", "Kill switch check before routing, risk validation, broker gateway delegation"],
        ["EndToEndFlowTest", "5 nested classes: KillSwitch, Settlement, OrderRecovery, OtpFlow, BrowserLock, CrossMarket"],
    ])

    doc.add_heading("Strategies", level=2)
    add_styled_table(doc, ["Class", "What It Tests"], [
        ["MomentumBreakoutStrategyTest", "Volume spike detection, RSI range filtering, SMA breakout signals"],
        ["EtfNavArbitrageStrategyTest", "NAV discount entry, premium exit, extreme premium handling"],
    ])

    doc.add_heading("Signals & Indicators", level=2)
    add_styled_table(doc, ["Class", "What It Tests"], [
        ["RsiCalculatorTest", "RSI(14) calculation accuracy, edge cases"],
        ["MacdCalculatorTest", "MACD line, signal line, histogram values"],
        ["NavDiscountCalculatorTest", "NAV discount/premium percentage calculation"],
    ])

    doc.add_heading("Notifications", level=2)
    add_styled_table(doc, ["Class", "What It Tests"], [
        ["MessageFormatterTest", "Trade signal formatting, portfolio summary formatting"],
        ["TradeApprovalServiceTest", "Approval flow, timeout behavior, default REJECT"],
    ])

    doc.add_heading("News", level=2)
    add_styled_table(doc, ["Class", "What It Tests"], [
        ["EventImpactRulesTest", "Event type to impact mapping"],
        ["NewsEventClassifierTest", "Headline classification into event types"],
    ])

    doc.add_heading("Long-term", level=2)
    add_styled_table(doc, ["Class", "What It Tests"], [
        ["DcaExecutorTest", "Monthly DCA execution logic, budget calculations"],
        ["DividendTrackerTest", "Ex-date tracking, alert generation"],
        ["PortfolioRebalancerTest", "Drift detection, rebalance action generation"],
    ])

    doc.add_heading("Backtest", level=2)
    add_styled_table(doc, ["Class", "What It Tests"], [
        ["PerformanceAnalyzerTest", "Sharpe ratio, max drawdown, win rate calculations"],
        ["SimulatedOrderExecutorTest", "Simulated fill logic with slippage"],
    ])

    doc.add_heading("Data & Discovery", level=2)
    add_styled_table(doc, ["Class", "What It Tests"], [
        ["EodhdApiClientTest", "API response parsing, OHLCV bar construction"],
        ["WatchlistManagerTest", "Watchlist add/remove, size limits, promotion/demotion"],
    ])

    # ── Integration Test Steps ──
    doc.add_heading("Integration Test Steps", level=1)

    steps = [
        ("Step 01: PostgreSQL + Flyway", "Step01_PostgresFlywayIT", "Database infrastructure", [
            ["1.1", "PostgreSQL connection alive", "Connection valid, product is PostgreSQL"],
            ["1.2", "Flyway ran all migrations", "flyway_schema_history >= 30 successful entries"],
            ["1.3", "All expected tables exist", "27 expected tables present in public schema"],
            ["1.4", "ohlcv_bars schema correct", "Columns: id, symbol, trade_date, open/high/low/close, volume"],
            ["1.5", "trade_orders schema correct", "Columns: id, order_id, symbol, side, quantity, price, status"],
            ["1.6", "No failed migrations", "Zero rows with success = false"],
            ["1.7", "Seed data exists", "watchlist_stocks has > 0 rows (from V20)"],
        ]),
        ("Step 02: EODHD Market Data", "Step02_EodhdApiIT", "Market data ingestion", [
            ["2.1", "API key configured", "EODHD_API_KEY is not blank"],
            ["2.2", "Fetch ZENITHBANK OHLCV", "Returns bars with positive close price and volume"],
            ["2.3", "Fetch GTCO OHLCV", "Returns recent bars with positive close"],
            ["2.4", "Fetch fundamentals", "Returns data (skipped on free-tier)"],
            ["2.5", "Screener returns candidates", "Returns NGX stocks (skipped on free tier)"],
            ["2.6", "OHLCV persists to DB", 'All bars have dataSource = "EODHD"'],
            ["2.7", "Values are reasonable", "High >= Low, Low <= Close <= High"],
        ]),
        ("Step 03: Telegram", "Step03_TelegramIT", "Telegram notifications", [
            ["3.1", "Credentials configured", "Bot token and chat ID not blank"],
            ["3.2", "Plain text message", "Sends without throwing"],
            ["3.3", "Formatted trade signal", "Rich formatted message sends"],
            ["3.4", "Long message", "Portfolio summary under 4096 char limit"],
        ]),
        ("Step 04: WhatsApp via WAHA", "Step04_WhatsAppIT", "WhatsApp notification + OTP", [
            ["4.1", "Chat ID configured", "Contains @c.us format"],
            ["4.2", "WAHA running", "/api/sessions returns response"],
            ["4.3", "Plain text message", "Sends without throwing"],
            ["4.4", "Formatted trade alert", "Sends trade signal format"],
            ["4.5", "OTP handler state", "No pending OTPs, max retries = 3"],
            ["4.6", "Webhook endpoint", "WhatsAppWebhookController bean loaded"],
        ]),
        ("Step 05: News Scrapers", "Step05_ScrapersIT", "News scraping from 6 sources", [
            ["5a", "BusinessDay", "Deduplicates by URL"],
            ["5b", "Nairametrics", "Same dedup behavior"],
            ["5c", "SeekingAlpha", "Often blocked (403 expected) — soft pass"],
            ["5d", "Reuters RSS", "Feed URL may change — soft pass"],
            ["5e", "CBN Press", "Website may change — soft pass"],
            ["5f", "NGX Price List + ASI Index", "May need JS rendering — soft pass"],
        ]),
        ("Step 06: Claude AI", "Step06_ClaudeAiIT", "AI analysis integration", [
            ["6.1", "AI configured", "AI_ENABLED=true, API key set"],
            ["6.2", "News analysis prompt", "Returns response with content + tokens"],
            ["6.3", "Cost tracker", "Records cost, daily/monthly totals"],
            ["6.4", "Budget check", "Budget not exceeded after test call"],
            ["6.5", "Structured analysis", "Response contains JSON-like structure"],
        ]),
        ("Step 07: Meritrade Login", "Step07_MeritradeLoginIT", "Browser automation login", [
            ["7.1", "Credentials configured", "Username and password not blank"],
            ["7.2", "Login succeeds", "Session active, screenshot exists"],
            ["7.3", "Portfolio holdings", "Returns Map (may be empty)"],
            ["7.4", "Available cash", "Returns BigDecimal >= 0"],
            ["7.5", "FX rate", "Returns USD/NGN rate (soft pass)"],
            ["7.6", "Dashboard screenshot", "Screenshot file exists on disk"],
        ]),
        ("Step 08: Paper Trade", "Step08_PaperTradeIT", "Real order submission (1 share)", [
            ["8.1", "Session active", "Broker session ready"],
            ["8.2", "Submit BUY LIMIT", "Order ID returned for TRANSCORP at NGN 5.00"],
            ["8.3", "Persist to DB", "TradeOrder saved with ID"],
            ["8.4", "Send notification", "Telegram notification sent"],
            ["8.5", "Check order status", "Status is not blank"],
            ["8.6", "Verify in DB", "Order found in database"],
        ]),
        ("Step 09: Backtesting", "Step09_BacktestIT", "Backtesting engine", [
            ["9.1", "Fetch historical data", "5 stocks, all have bars"],
            ["9.2", "List strategies", "At least 1 strategy registered"],
            ["9.3", "Run backtest", "BacktestRun returned with metrics"],
            ["9.4", "Results persist", "Rows in backtest_runs and backtest_trades"],
            ["9.5", "Metrics reasonable", "Max drawdown <= 0, win rate 0-100%"],
        ]),
        ("Step 10: Dashboard REST API", "Step10_DashboardIT", "REST API endpoints", [
            ["10.1-14", "All 14 endpoints return 200 OK", "/api/portfolio, /fx, /settlement, /performance, /signals, /news, /ai/cost, /discovery/active, /discovery/candidates, /backtest/runs, /backtest/strategies, /actuator/health, /killswitch, /dividends"],
        ]),
        ("Step 11: End-to-End Pipeline", "Step11_EndToEndIT", "Full signal-to-notification pipeline", [
            ["E2E 1/10", "Fetch 3 months OHLCV", "ZENITHBANK data loaded"],
            ["E2E 2/10", "Calculate indicators", "RSI(14) and MACD computed"],
            ["E2E 3/10", "Scrape news", "BusinessDay articles fetched"],
            ["E2E 4/10", "AI analysis", "News + market data analyzed"],
            ["E2E 5/10", "Generate signal", "Mock BUY signal created"],
            ["E2E 6/10", "Safety checks", "Kill switch OFF, circuit breakers OK"],
            ["E2E 7/10", "Approval request", "Sent via Telegram"],
            ["E2E 8/10", "Simulate approval", "Auto-approved"],
            ["E2E 9/10", "Record trade", "Saved to PostgreSQL"],
            ["E2E 10/10", "Send confirmation", "Telegram + WhatsApp notified"],
        ]),
    ]

    for title, class_name, desc, tests in steps:
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Class: {class_name} — Validates: {desc}")
        add_styled_table(doc, ["Test", "Assertion", "Details"],
                         [[t[0], t[1], t[2]] for t in tests], [1.0, 2.5, 3.0])

    # ── Running Tests ──
    doc.add_heading("Running Tests", level=1)
    add_styled_table(doc, ["Command", "What Runs"], [
        ["mvn test", "Unit tests only (excludes integration)"],
        ["mvn verify -Pintegration", "Integration tests only"],
        ['mvn verify -Pintegration -Dtest="Step01_PostgresFlywayIT"', "Single integration step"],
    ])

    # ── Manual Testing Checklist ──
    doc.add_heading("Manual Testing Checklist", level=1)

    doc.add_heading("Login Flow", level=2)
    for item in [
        "Bot navigates to https://app.trovefinance.com/login",
        "Email and password fields are filled",
        "Login button is clicked",
        "If OTP screen appears, bot sends WhatsApp/email prompt",
        "OTP is entered and verified",
        "Dashboard loads (verify via screenshot in ./screenshots/)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)

    doc.add_heading("Kill Switch", level=2)
    for item in [
        "Activate kill switch via KillSwitchService.activate(reason)",
        "Verify all trading halts immediately",
        "Verify urgent WhatsApp/Telegram notification sent",
        "Deactivate and verify trading resumes",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)

    # ── Known Limitations ──
    doc.add_heading("Known Limitations", level=1)
    limitations = [
        "Portfolio selectors need real account with positions to validate",
        "submitOrder() not yet fully wired for production",
        "Trove requires account verification before trading is enabled",
        "SeekingAlpha scraper frequently blocked by anti-bot (403 expected)",
        "RSS feed URLs may change without notice",
        "EODHD free tier: fundamentals/screener require paid plan",
    ]
    for i, lim in enumerate(limitations, 1):
        doc.add_paragraph(f"{i}. {lim}")

    save_doc(doc, "QA_GUIDE.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 3: BUSINESS OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

def generate_business_overview():
    print("\n[3/7] Generating BUSINESS_OVERVIEW.docx...")
    doc = Document()
    setup_styles(doc)
    add_title_page(doc, "Business Overview", "Executive Summary & Market Position",
                   "Business stakeholders, management, and non-technical decision makers")

    # ── Executive Summary ──
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "The NGX Trading Bot is an autonomous software system that trades stocks on the Nigerian Stock Exchange (NGX) "
        "and US equity markets through the Trove/Meritrade brokerage platform. It eliminates the need for manual broker "
        "interaction by using browser automation technology to log in, place orders, and manage a portfolio — just as a "
        "human trader would, but faster, more disciplined, and available during every second of market hours."
    )

    # ── Problem Statement ──
    doc.add_heading("Problem Statement", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Nigerian brokers provide no trading APIs.")
    run.font.bold = True
    p.add_run(
        " Unlike developed markets where algorithmic trading is standard, the NGX ecosystem requires traders to "
        "manually log into web portals, navigate dashboards, and click through multi-step order forms."
    )

    add_styled_table(doc, ["Challenge", "Impact"], [
        ["Speed", "Manual order entry takes minutes. In a 4.5-hour window (10:00-14:30 WAT), delays cost money."],
        ["Discipline", "Human traders deviate from strategy under pressure — chasing losses, holding too long."],
        ["Coverage", "A single person cannot monitor 50+ stocks, 6 news sources, and execute orders simultaneously."],
        ["Errors", "Manual data entry leads to wrong quantities, wrong prices, and missed opportunities."],
    ])

    # ── Solution ──
    doc.add_heading("The Solution", level=1)
    doc.add_paragraph("The NGX Trading Bot acts as an autonomous digital trader that sees, thinks, decides, acts, and reports:")

    add_diagram(doc, """graph LR
    subgraph SEES["1. SEES"]
        A1["Market Data\\n(EODHD API)"]
        A2["Financial News\\n(6 sources)"]
        A3["Broker Portfolio\\n(Trove dashboard)"]
    end

    subgraph THINKS["2. THINKS"]
        B1["10 Trading\\nStrategies"]
        B2["Technical\\nIndicators"]
        B3["AI Analysis\\n(Claude)"]
    end

    subgraph DECIDES["3. DECIDES"]
        C1["7 Risk Rules"]
        C2["Circuit\\nBreakers"]
        C3["Kill Switch\\nCheck"]
    end

    subgraph ACTS["4. ACTS"]
        D1["Browser\\nAutomation"]
        D2["Auto-Login\\n+ OTP"]
        D3["LIMIT Order\\nPlacement"]
    end

    subgraph REPORTS["5. REPORTS"]
        E1["WhatsApp\\nAlerts"]
        E2["Telegram\\nAlerts"]
        E3["Approval\\nRequests"]
    end

    SEES --> THINKS --> DECIDES --> ACTS --> REPORTS
    """, "Figure 1: Solution Flow — Sees, Thinks, Decides, Acts, Reports", 6.0)

    # ── Key Capabilities ──
    doc.add_heading("Key Capabilities", level=1)

    doc.add_heading("Autonomous Broker Interaction", level=2)
    for item in [
        "Logs into Trove/Meritrade without human intervention",
        "Handles OTP verification automatically via Gmail or WhatsApp",
        "Reads portfolio holdings and cash balances from broker dashboard",
        "Places LIMIT orders only (never market orders, for price protection)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)

    doc.add_heading("10 Trading Strategies", level=2)
    add_styled_table(doc, ["Strategy", "Focus"], [
        ["Momentum Breakout", "Volume spikes + price breakouts on NGX"],
        ["ETF NAV Arbitrage", "Buy ETFs trading below net asset value"],
        ["Dollar-Cost Averaging", "Monthly systematic buying"],
        ["Dividend Accumulation", "High-yield NGX and US stocks"],
        ["Value Accumulation", "Fundamentally undervalued NGX stocks"],
        ["Sector Rotation", "Rotate into strongest NGX sectors"],
        ["Pension Flow Overlay", "Follow pension fund capital flows"],
        ["US Earnings Momentum", "Trade post-earnings moves on US stocks"],
        ["US ETF Rotation", "Monthly sector ETF rotation (US)"],
        ["Currency Hedge", "Gold/USD positions during Naira weakness"],
    ])

    # ── Risk Management ──
    doc.add_heading("Risk Management", level=1)
    doc.add_paragraph("Every trade must pass ALL of these checks — no exceptions:")

    add_styled_table(doc, ["Rule", "Limit", "Purpose"], [
        ["Max risk per trade", "2% of portfolio", "No single trade can cause large losses"],
        ["Max single position", "15% (20% for core)", "Prevents overconcentration"],
        ["Max sector exposure", "40% of portfolio", "Diversification across sectors"],
        ["Min cash reserve", "20% of portfolio", "Always maintain liquidity"],
        ["Max open positions", "10", "Manageable position count"],
        ["Min avg daily volume", "10,000 shares", "Ensures stock is liquid"],
        ["Max volume participation", "10% of daily volume", "Avoids market impact"],
    ])

    doc.add_heading("Safety Layers", level=2)
    add_diagram(doc, """graph TB
    SIGNAL["Trade Signal Generated"] --> R1

    subgraph SAFETY["Safety Gates (ALL must pass)"]
        R1["Gate 1: Kill Switch OFF?"] --> R2
        R2["Gate 2: Circuit Breaker OK?"] --> R3
        R3["Gate 3: Max Positions < 10?"] --> R4
        R4["Gate 4: Position Size < 15%?"] --> R5
        R5["Gate 5: Sector Exposure < 40%?"] --> R6
        R6["Gate 6: Cash Reserve > 20%?"] --> R7
        R7["Gate 7: Risk Per Trade < 2%?"] --> APPROVE
    end

    APPROVE["Human Approval\\n(WhatsApp/Telegram)"] --> EXEC["Order Executed"]
    R1 -->|FAIL| BLOCK["Trade BLOCKED"]
    R2 -->|FAIL| BLOCK
    R3 -->|FAIL| BLOCK
    R4 -->|FAIL| BLOCK
    R5 -->|FAIL| BLOCK
    R6 -->|FAIL| BLOCK
    R7 -->|FAIL| BLOCK

    style BLOCK fill:#f03e3e,color:#fff
    style EXEC fill:#51cf66,color:#fff
    style APPROVE fill:#ffa94d,color:#fff
    """, "Figure 2: 7-Gate Safety System", 5.0)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Circuit Breakers: ")
    run.font.bold = True
    p.add_run("5% daily loss halts trading for the day. 10% weekly loss halts until Monday.")

    p = doc.add_paragraph()
    run = p.add_run("Kill Switch: ")
    run.font.bold = True
    p.add_run("Instant manual override to halt all trading with one command. Sends urgent notification.")

    # ── Market Coverage ──
    doc.add_heading("Market Coverage", level=1)

    doc.add_heading("NGX (Nigerian Stock Exchange)", level=2)
    doc.add_paragraph(
        "Large caps: Zenith Bank, GTCO, Access Corp, UBA, Dangote Cement, Seplat, MTNN. "
        "ETFs: Stanbic ETF30, VetGrif30, MerGrowth, MerValue, SiamletETF40, NewGold. "
        "Market hours: 10:00-14:30 WAT. Settlement: T+2. Daily price limit: +/-10%."
    )

    doc.add_heading("US Equities (via Trove)", level=2)
    doc.add_paragraph(
        "ETFs: VOO, SCHD, BND, GLD, VXUS, VNQ. Individual US stocks for earnings momentum. Settlement: T+1."
    )

    # ── Competitive Advantages ──
    doc.add_heading("Competitive Advantages", level=1)
    advantages = [
        ("First-mover in NGX automation", "No known competitors offer autonomous trading for the Nigerian market."),
        ("No API dependency", "Browser automation works regardless of broker API availability."),
        ("AI-enhanced analysis", "Claude AI for sophisticated news and earnings analysis."),
        ("Settlement-aware cash management", "Tracks settled vs. unsettled cash (T+2 NGX, T+1 US)."),
        ("Dual-market support", "Trades both NGX and US equities through a single platform."),
        ("News intelligence", "Scrapes 6 sources and classifies events by market impact."),
    ]
    for title, desc in advantages:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        p.add_run(desc)

    # ── Current Status ──
    doc.add_heading("Current Status", level=1)
    add_styled_table(doc, ["Phase", "Status", "Details"], [
        ["Phase 1: Infrastructure", "Complete", "Autonomous login, OTP, 10 strategies, risk engine, notifications, AI, news, backtesting, 30 migrations, 181+ tests"],
        ["Phase 2: Live Trading", "Next", "Live order execution, portfolio sync, P&L tracking, account verification"],
        ["Phase 3: Scale", "Future", "Multi-broker support, mobile dashboard, walk-forward optimization"],
    ])

    save_doc(doc, "BUSINESS_OVERVIEW.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 4: PRODUCT SPEC
# ═══════════════════════════════════════════════════════════════════════════════

def generate_product_spec():
    print("\n[4/7] Generating PRODUCT_SPEC.docx...")
    doc = Document()
    setup_styles(doc)
    add_title_page(doc, "Product Specification", "Feature Inventory & Roadmap",
                   "Product managers, designers, and project planners")

    # ── Product Vision ──
    doc.add_heading("Product Vision", level=1)
    doc.add_paragraph(
        "A fully autonomous trading system for emerging market equities, starting with the Nigerian Stock Exchange (NGX) "
        "and expanding to US equities via Trove's dual-market platform. The system replaces manual broker interaction with "
        "intelligent automation — analyzing data, generating signals, managing risk, executing trades, and reporting results "
        "with zero human intervention during normal operation."
    )

    # ── Feature Inventory ──
    doc.add_heading("Feature Inventory", level=1)

    doc.add_heading("Core Trading Engine", level=2)
    add_styled_table(doc, ["Feature", "Status", "Description"], [
        ["Signal generation pipeline", "Done", "Technical + fundamental indicators scored into signals"],
        ["Technical indicators", "Done", "RSI, MACD, ATR, SMA/EMA, volume analysis"],
        ["Fundamental signals", "Done", "NAV discount, dividend proximity, PenCom eligibility"],
        ["Composite signal scoring", "Done", "Weighted multi-indicator scoring"],
        ["Strategy evaluation", "Done", "10 strategies with CORE/SATELLITE pool classification"],
        ["Strategy enable/disable", "Done", "Per-strategy YAML toggle"],
        ["Market hours enforcement", "Done", "Hard-coded 10:00-14:30 WAT check"],
        ["Order routing", "Done", "Signal -> risk check -> approval -> execution pipeline"],
    ])

    doc.add_heading("Trading Strategies", level=2)
    add_styled_table(doc, ["Strategy", "Status", "Market", "Pool"], [
        ["Momentum Breakout", "Done", "NGX", "SATELLITE"],
        ["ETF NAV Arbitrage", "Done", "NGX", "SATELLITE"],
        ["Dollar-Cost Averaging", "Done", "BOTH", "CORE"],
        ["Dividend Accumulation", "Done", "BOTH", "CORE"],
        ["Value Accumulation", "Done", "NGX", "CORE"],
        ["Sector Rotation", "Done", "NGX", "SATELLITE"],
        ["Pension Flow Overlay", "Done", "NGX", "SATELLITE"],
        ["US Earnings Momentum", "Done", "US", "SATELLITE"],
        ["US ETF Rotation", "Done", "US", "SATELLITE"],
        ["Currency Hedge", "Done", "BOTH", "CORE"],
    ])

    doc.add_heading("Broker Integration", level=2)
    add_styled_table(doc, ["Feature", "Status", "Description"], [
        ["Playwright browser automation", "Done", "Chromium-based browser control"],
        ["Trove login automation", "Done", "Navigate to login, fill credentials, submit"],
        ["OTP via WhatsApp", "Done", "Bot requests OTP, user replies via WhatsApp"],
        ["OTP via Gmail IMAP", "Done", "Auto-reads OTP from Trove emails"],
        ["Portfolio scraping", "Done", "Reads holdings and quantities from broker"],
        ["Cash balance reading", "Done", "Reads available cash from broker"],
        ["FX rate reading", "Done", "Reads USD/NGN rate from broker"],
        ["Screenshot capture", "Done", "Saves browser state at key moments"],
        ["Order submission", "In Progress", "submitOrder() implemented, pending live wiring"],
        ["Order status checking", "Done", "Reads order status from broker"],
        ["Browser session management", "Done", "Session lock, max hours, auto-relogin"],
        ["CSS selector configuration", "Done", "All UI selectors configurable via YAML"],
    ])

    doc.add_heading("Risk Management", level=2)
    add_styled_table(doc, ["Feature", "Status", "Description"], [
        ["Risk per trade (2%)", "Done", "Validates max portfolio risk per trade"],
        ["Single position (15%/20%)", "Done", "Pool-aware: 15% satellite, 20% core"],
        ["Sector exposure (40%)", "Done", "Cross-position sector aggregation"],
        ["Cash reserve (20%)", "Done", "Ensures liquidity"],
        ["Max open positions (10)", "Done", "Hard cap on concurrent positions"],
        ["Volume participation (10%)", "Done", "Prevents market impact"],
        ["Daily circuit breaker (5%)", "Done", "Halts on 5% daily loss"],
        ["Weekly circuit breaker (10%)", "Done", "Halts until Monday on 10% weekly loss"],
        ["Kill switch", "Done", "Manual emergency halt with notification"],
        ["Settlement cash tracking", "Done", "T+2 NGX, T+1 US"],
        ["Order recovery", "Done", "UNCERTAIN status + kill switch on failure"],
        ["Stop-loss monitoring", "Done", "Monitors positions against stop levels"],
        ["Position sizing", "Done", "Optimal quantity given risk constraints"],
    ])

    doc.add_heading("Notifications", level=2)
    add_styled_table(doc, ["Feature", "Status", "Description"], [
        ["WhatsApp via WAHA", "Done", "Primary channel"],
        ["Telegram Bot API", "Done", "Fallback channel with rich formatting"],
        ["Message formatting", "Done", "Structured signal/portfolio/alert templates"],
        ["Trade approval flow", "Done", "Human-in-the-loop with 5-min timeout"],
        ["OTP request/response", "Done", "Bidirectional WhatsApp for OTP"],
        ["Webhook receiver", "Done", "WAHA POSTs to /api/webhooks/whatsapp/message"],
        ["Urgent notifications", "Done", "Kill switch, circuit breaker, failures"],
    ])

    doc.add_heading("AI Integration", level=2)
    add_styled_table(doc, ["Feature", "Status", "Description"], [
        ["Claude API client", "Done", "Sends prompts to Anthropic API"],
        ["News sentiment analysis", "Done", "BULLISH/BEARISH/NEUTRAL with confidence"],
        ["Earnings analysis", "Done", "EPS/revenue surprise assessment"],
        ["Cross-article synthesis", "Done", "Correlates themes across sources"],
        ["Cost tracking", "Done", "Per-call cost + daily/monthly totals"],
        ["Budget enforcement", "Done", "$5/day, $100/month limits"],
        ["Deep analysis triggers", "Done", "Auto-upgrade to stronger model for key events"],
        ["Graceful fallback", "Done", "Continues without AI when unavailable"],
    ])

    doc.add_heading("News Intelligence", level=2)
    add_styled_table(doc, ["Feature", "Status", "Description"], [
        ["BusinessDay scraper", "Done", "Nigerian business news"],
        ["Nairametrics scraper", "Done", "Nigerian financial news"],
        ["SeekingAlpha scraper", "Done", "US market analysis (may be blocked)"],
        ["Reuters RSS", "Done", "International financial RSS"],
        ["CBN Press scraper", "Done", "Central Bank of Nigeria releases"],
        ["NGX Bulletin parser", "Done", "Exchange announcements"],
        ["Event classification", "Done", "Categorizes by event type"],
        ["Impact scoring", "Done", "Rates by market impact severity"],
        ["URL deduplication", "Done", "Prevents duplicate processing"],
    ])

    doc.add_heading("Stock Discovery", level=2)
    add_styled_table(doc, ["Feature", "Status", "Description"], [
        ["EODHD screener", "Done", "Screens NGX stocks by fundamentals"],
        ["Watchlist management", "Done", "Max 30 active, max 20 observation"],
        ["Candidate evaluation", "Done", "Fundamental scoring"],
        ["Promotion/demotion policies", "Done", "Rules for tier movement"],
        ["Observation period", "Done", "14-day evaluation"],
        ["Demotion cooldown", "Done", "90-day cooldown"],
    ])

    doc.add_heading("Long-Term Portfolio", level=2)
    add_styled_table(doc, ["Feature", "Status", "Description"], [
        ["DCA scheduling", "Done", "Monthly: NGX day 5, US day 10"],
        ["DCA budgets", "Done", "NGN 150K/month NGX, $300/month US"],
        ["Dividend tracking", "Done", "Ex-date monitoring with 7-day alerts"],
        ["Dividend reinvestment", "Done", "Auto-reinvest into same stock"],
        ["Portfolio rebalancing", "Done", "Quarterly with 10% drift threshold"],
        ["Target allocation", "Done", "14 positions across NGX + US"],
    ])

    doc.add_heading("Backtesting", level=2)
    add_styled_table(doc, ["Feature", "Status", "Description"], [
        ["Backtest runner", "Done", "Strategies against historical data"],
        ["Simulated executor", "Done", "Fill simulation with slippage"],
        ["Performance analyzer", "Done", "Sharpe, drawdown, win rate"],
        ["Equity curve", "Done", "Point-by-point tracking"],
        ["Results persistence", "Done", "Stored in PostgreSQL"],
        ["REST API", "Done", "Trigger and view via API"],
    ])

    # ── User Flows ──
    doc.add_heading("User Flows", level=1)

    doc.add_heading("Signal-to-Execution Flow", level=2)
    add_diagram(doc, """flowchart TD
    START(["Market Open\\n10:00 WAT"]) --> DATA

    subgraph DATA["1. Data Collection"]
        D1["EODHD OHLCV"] & D2["News Scrapers (6)"] & D3["NGX Web Scraper"]
    end

    DATA --> SIGNAL

    subgraph SIGNAL["2. Signal Generation"]
        S1["RSI, MACD, ATR, SMA"] --> S2["CompositeSignalScorer"]
        S3["NAV Discount, Dividends"] --> S2
    end

    SIGNAL --> STRATEGY

    subgraph STRATEGY["3. Strategy Evaluation"]
        ST1["10 Strategies evaluate()"]
        ST1 --> ST2{"Signal\\ngenerated?"}
    end

    ST2 -->|No| DONE(["No action"])
    ST2 -->|Yes| RISK

    subgraph RISK["4. Risk Validation"]
        R1["KillSwitch check"] --> R2["CircuitBreaker check"] --> R3["RiskManager: 5 checks"]
    end

    R3 --> R4{"All\\npassed?"}
    R4 -->|No| BLOCKED["Trade BLOCKED"]
    R4 -->|Yes| APPROVAL

    subgraph APPROVAL["5. Trade Approval"]
        A1["WhatsApp/Telegram\\napproval request"] --> A2{"User replies\\nwithin 5 min?"}
    end

    A2 -->|No / REJECT| REJECTED["Auto-REJECTED"]
    A2 -->|YES| EXEC

    subgraph EXEC["6. Order Execution"]
        E1["BrowserSession.acquire()"] --> E2["TroveBrowserAgent.login()"] --> E5["submitOrder(LIMIT)"]
    end

    E5 --> NOTIFY["7. Notification\\nWhatsApp + Telegram"]

    style BLOCKED fill:#f03e3e,color:#fff
    style REJECTED fill:#f03e3e,color:#fff
    style DONE fill:#868e96,color:#fff
    """, "Figure 1: Signal-to-Execution Pipeline", 5.5)

    doc.add_heading("Order State Machine", level=2)
    add_diagram(doc, """stateDiagram-v2
    [*] --> PENDING : Signal generated
    PENDING --> SUBMITTED : Sent to broker
    SUBMITTED --> FILLED : Broker confirms fill
    SUBMITTED --> PARTIAL_FILL : Partial execution
    PARTIAL_FILL --> FILLED : Remaining filled
    SUBMITTED --> CANCELLED : User/system cancels
    SUBMITTED --> REJECTED : Broker rejects
    SUBMITTED --> UNCERTAIN : Execution failure
    UNCERTAIN --> FILLED : Manual verification
    UNCERTAIN --> CANCELLED : Manual cancellation
    PENDING --> CANCELLED : Risk check fails
    FILLED --> [*]
    CANCELLED --> [*]
    REJECTED --> [*]
    """, "Figure 2: Order State Machine", 5.0)

    doc.add_heading("Watchlist State Machine", level=2)
    add_diagram(doc, """stateDiagram-v2
    [*] --> DISCOVERED : Screener finds stock
    DISCOVERED --> OBSERVATION : Meets minimum criteria
    OBSERVATION --> ACTIVE : 14-day evaluation passes
    OBSERVATION --> EXCLUDED : Fails evaluation
    ACTIVE --> DEMOTED : No signals for 60 days
    DEMOTED --> OBSERVATION : After 90-day cooldown
    DEMOTED --> EXCLUDED : Fundamentals deteriorate
    ACTIVE --> EXCLUDED : Delisted or illiquid
    """, "Figure 3: Watchlist State Machine", 5.0)

    # ── Data Model ──
    doc.add_heading("Data Model", level=1)
    add_diagram(doc, """erDiagram
    watchlist_stocks ||--o{ ohlcv_bars : "has price data"
    ohlcv_bars ||--o{ trade_signals : "generates"
    trade_signals ||--o{ trade_orders : "triggers"
    trade_orders ||--o{ positions : "creates"
    trade_orders ||--o{ approval_requests : "requires"
    trade_orders ||--o{ notification_log : "notifies about"
    positions }o--|| portfolio_snapshots : "aggregated in"
    news_items ||--o{ ai_analysis : "analyzed by"
    ai_analysis ||--o{ ai_cost_ledger : "costs tracked in"
    discovered_stocks ||--o{ discovery_events : "audit trail"
    discovered_stocks ||--o{ watchlist_stocks : "promoted to"
    dca_plans ||--o{ trade_orders : "executes via"
    dividend_events ||--o{ trade_orders : "reinvests via"
    backtest_runs ||--o{ backtest_trades : "contains"
    backtest_runs ||--o{ equity_curve_points : "tracks equity"
    """, "Figure 4: Entity Relationship Diagram", 5.5)

    # ── Roadmap ──
    doc.add_heading("Roadmap", level=1)
    add_styled_table(doc, ["Phase", "Status", "Key Deliverables"], [
        ["Phase 1: Infrastructure & Intelligence", "Complete", "30 migrations, 10 strategies, risk engine, browser automation, OTP, notifications, AI, news, backtesting, 181+ tests"],
        ["Phase 2: Live Trading", "Next", "Wire submitOrder() for production, real-time portfolio sync, P&L tracking, Trove account verification"],
        ["Phase 3: Scale & Optimize", "Future", "Multi-broker support, mobile dashboard, walk-forward optimization, ML signal enhancement"],
    ])

    # ── Success Metrics ──
    doc.add_heading("Success Metrics", level=1)
    add_styled_table(doc, ["Metric", "Target", "Measurement"], [
        ["OTP success rate", "> 95%", "OTP attempts vs. successful verifications"],
        ["Risk rule compliance", "100%", "Zero trades with failed risk checks"],
        ["Notification latency", "< 5 seconds", "Signal to WhatsApp/Telegram delivery"],
        ["Unit test pass rate", "100%", "mvn test green on every commit"],
        ["Integration test pass rate", "> 90%", "Steps 01-11 (some depend on external services)"],
        ["Kill switch response", "< 1 second", "Activation to trading halt"],
        ["Backtest Sharpe", "> 0.5", "1-year historical backtest"],
        ["Uptime (market hours)", "> 99%", "10:00-14:30 WAT availability"],
    ])

    save_doc(doc, "PRODUCT_SPEC.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 5: PITCH
# ═══════════════════════════════════════════════════════════════════════════════

def generate_pitch():
    print("\n[5/7] Generating PITCH.docx...")
    doc = Document()
    setup_styles(doc)
    add_title_page(doc, "NGX Trading Bot", "Investor Pitch",
                   "Investors, accelerators, and pitch competitions")

    # ── The Problem ──
    doc.add_heading("The Problem", level=1)
    doc.add_paragraph(
        "Nigeria's stock market (NGX) has a market capitalization of over $50 billion and is one of Africa's largest "
        "exchanges. Yet retail investors face a frustrating reality:"
    )
    problems = [
        ("Zero API access", "Nigerian brokers offer no programmatic trading APIs. Every order must be placed manually."),
        ("4.5-hour trading window", "The NGX is open 10:00-14:30 WAT. Miss a signal, miss the opportunity."),
        ("Manual, error-prone execution", "10+ manual steps per trade: login, navigate, enter details, handle OTP."),
        ("No automation tools exist", "While developed markets have dozens of algo platforms, Nigeria has none."),
        ("Growing retail participation", "NGX retail accounts are growing, but tools remain primitive."),
    ]
    for title, desc in problems:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title} — ")
        run.font.bold = True
        p.add_run(desc)

    # ── The Solution ──
    doc.add_heading("The Solution", level=1)

    doc.add_heading("Manual Trading vs. NGX Trading Bot", level=2)
    add_diagram(doc, """graph LR
    subgraph BEFORE["Manual Trading (Today)"]
        direction TB
        M1["Open broker website"] --> M2["Type username/password"]
        M2 --> M3["Enter OTP from phone"]
        M3 --> M4["Navigate to trade screen"]
        M4 --> M5["Search for stock"]
        M5 --> M6["Enter quantity & price"]
        M6 --> M7["Click review, then confirm"]
        M7 --> M8["Hope you didn't miss the price move"]
    end

    subgraph AFTER["NGX Trading Bot (Automated)"]
        direction TB
        A1["Bot monitors 50+ stocks 24/7"] --> A2["AI detects opportunity"]
        A2 --> A3["Risk engine validates"]
        A3 --> A4["You approve on WhatsApp"]
        A4 --> A5["Order executed in seconds"]
    end

    style BEFORE fill:#f8d7da,stroke:#f03e3e
    style AFTER fill:#d3f9d8,stroke:#51cf66
    """, "Before vs. After: Manual Trading vs. NGX Trading Bot", 6.0)

    p = doc.add_paragraph()
    run = p.add_run(
        "An AI-powered autonomous trading bot that sees, thinks, and acts like a human trader — "
        "but faster, more disciplined, and never misses market hours."
    )
    run.font.bold = True
    run.font.size = Pt(12)

    # ── How It Works ──
    doc.add_heading("How It Works", level=1)
    add_styled_table(doc, ["Stage", "Components", "Purpose"], [
        ["1. DATA", "EODHD Market Data, 6 News Sources, Broker Portfolio", "Collect real-time information"],
        ["2. ANALYSIS", "RSI, MACD, ATR + AI Sentiment + Fundamental Screens", "Generate insights from data"],
        ["3. DECISION", "10 Strategies + Signal Scoring + Risk Validation", "Determine optimal trade actions"],
        ["4. EXECUTION", "Browser Automation (Playwright) + Auto-Login + Auto-OTP", "Execute LIMIT orders"],
        ["5. ALERT", "WhatsApp + Telegram + Human Approval for Big Trades", "Keep you informed and in control"],
    ])

    # ── Market Opportunity ──
    doc.add_heading("Market Opportunity", level=1)
    opportunities = [
        ("NGX market cap", "~$50B+ with growing daily trading volumes"),
        ("Retail growth", "Increasing year-over-year, driven by fintech awareness"),
        ("Automation gap", "Zero existing solutions for automated NGX trading"),
        ("Dual-market access", "Trove enables both NGX and US equities from one platform"),
        ("Emerging market alpha", "Less efficient markets offer greater systematic opportunities"),
        ("Regulatory tailwind", "NGX actively modernizing infrastructure"),
    ]
    for title, desc in opportunities:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{title}: ")
        run.font.bold = True
        p.add_run(desc)

    # ── Technology Moat ──
    doc.add_heading("Technology Moat", level=1)

    moats = [
        ("Browser Automation Bypasses API Absence",
         "Where others see a blocker (no APIs), we see a solved problem. Playwright-based automation mimics human "
         "interaction, working with any broker UI regardless of API availability."),
        ("Zero-Touch OTP Handling",
         "Our bot reads OTP codes from Gmail (IMAP) and WhatsApp (WAHA), entering them automatically. "
         "No human needs to be present to verify login."),
        ("AI-Enhanced Signal Analysis",
         "Claude AI analyzes news from 6 sources, interprets earnings reports, and detects insider trading patterns. "
         "Automatic escalation to deeper analysis for high-impact events."),
        ("Settlement-Aware Cash Management",
         "Independently tracks settled vs. unsettled cash across both markets (T+2 NGX, T+1 US), "
         "preventing the bot from spending money that hasn't settled."),
        ("Configurable Without Code Changes",
         "Every strategy, risk parameter, and UI selector is configurable via YAML. "
         "Broker UI changes = config file update, not code deployment."),
    ]
    for i, (title, desc) in enumerate(moats, 1):
        doc.add_heading(f"{i}. {title}", level=2)
        doc.add_paragraph(desc)

    # ── Traction ──
    doc.add_heading("Traction", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Working prototype with proven infrastructure:")
    run.font.bold = True

    traction = [
        "Autonomous login + OTP handling (Gmail + WhatsApp)",
        "10 trading strategies implemented and tested",
        "7-rule risk management engine with circuit breakers",
        "Real-time WhatsApp and Telegram notifications",
        "AI analysis integration with budget controls",
        "News intelligence from 6 sources",
        "Backtesting engine with historical data",
        "Dashboard REST API (29 endpoints)",
        "181+ unit tests passing",
        "11 integration test steps validating the full pipeline",
        "30 database migrations for a production-ready schema",
        "Dual-market support: NGX and US equities",
    ]
    for item in traction:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)

    # ── Revenue Model ──
    doc.add_heading("Revenue Model", level=1)
    add_styled_table(doc, ["Model", "Description", "Target"], [
        ["SaaS Subscription", "Monthly/annual access to the trading bot platform", "Retail investors, HNWIs"],
        ["Performance Fees", "Percentage of profits above a benchmark", "Managed account clients"],
        ["White-Label", "Licensed to Nigerian brokerages as a value-add", "Meristem, Stanbic IBTC, CSCS participants"],
        ["Data & Analytics", "NGX market intelligence, sentiment, signal feeds", "Institutional investors, asset managers"],
    ])

    # ── Team ──
    doc.add_heading("Team", level=1)
    add_callout(doc, "[Placeholder — Add team member bios, relevant experience, and LinkedIn profiles]")

    # ── The Ask ──
    doc.add_heading("The Ask", level=1)
    add_callout(doc, "[Placeholder — Customize based on your fundraising stage and needs]")
    doc.add_paragraph("Potential use of funds:")
    for item in [
        "Complete live trading integration and deploy to production",
        "Onboard first 100 beta users with managed accounts",
        "Build mobile dashboard for portfolio monitoring",
        "Expand to additional Nigerian brokerages",
        "Hire quantitative analyst and full-stack engineer",
    ]:
        doc.add_paragraph(f"  {item}", style="List Bullet")

    # ── Contact ──
    doc.add_heading("Contact", level=1)
    add_callout(doc, "[Placeholder — Add your email, phone, website, and GitHub repository link]")

    save_doc(doc, "PITCH.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 6: API REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════

def generate_api_reference():
    print("\n[6/7] Generating API_REFERENCE.docx...")
    doc = Document()
    setup_styles(doc)
    add_title_page(doc, "API Reference", "REST API Documentation",
                   "Frontend developers, integration engineers, and API consumers")

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "The NGX Trading Bot exposes a REST API on port 8080 for dashboard consumption, backtest management, "
        "and system control. All responses are JSON. No authentication is currently required "
        "(designed for local/private network access)."
    )
    doc.add_paragraph("Base URL: http://localhost:8080")

    doc.add_heading("Endpoint Map", level=2)
    add_diagram(doc, """graph LR
    subgraph PORTFOLIO["Portfolio & Dashboard"]
        P1["GET /api/portfolio"]
        P2["GET /api/portfolio/ngx"]
        P3["GET /api/portfolio/us"]
        P4["GET /api/portfolio/core"]
        P5["GET /api/portfolio/satellite"]
        P6["GET /api/fx"]
        P7["GET /api/settlement"]
        P8["GET /api/performance"]
        P9["GET /api/dividends"]
    end

    subgraph SIGNALS["Signals & Intelligence"]
        S1["GET /api/signals"]
        S2["GET /api/news"]
        S3["GET /api/ai/cost"]
        S4["GET /api/ai/analysis/{symbol}"]
        S5["GET /api/ai/fallback-rate"]
    end

    subgraph BACKTEST["Backtesting"]
        B1["POST /api/backtest/run"]
        B2["GET /api/backtest/runs"]
        B3["GET /api/backtest/runs/{id}"]
        B4["GET /api/backtest/strategies"]
    end

    subgraph CONTROL["System Control"]
        K1["GET /api/killswitch"]
        K2["POST /api/killswitch/activate"]
        K3["POST /api/killswitch/deactivate"]
    end
    """, "Figure 1: Endpoint Map", 6.0)

    # ── Portfolio Endpoints ──
    doc.add_heading("Portfolio & Dashboard", level=1)

    endpoints = [
        ("GET /api/portfolio", "Returns consolidated portfolio overview across both markets.",
         '{\n  "totalValueNgn": 5250000.00,\n  "ngxValue": 3750000.00,\n  "usValueUsd": 1200.00,\n  "fxRate": 1250.00,\n  "openPositions": 7,\n  "byMarket": { "NGX": 3750000.00, "US": 1200.00 },\n  "byPool": { "CORE": 4000000.00, "SATELLITE": 1250000.00 },\n  "killSwitchActive": false\n}'),
        ("GET /api/portfolio/ngx", "Returns NGX-only portfolio with positions and available cash.",
         '{\n  "market": "NGX",\n  "currency": "NGN",\n  "totalValue": 3750000.00,\n  "positions": [...],\n  "positionCount": 5,\n  "availableCash": 850000.00\n}'),
        ("GET /api/portfolio/us", "Returns US-only portfolio.",
         '{\n  "market": "US",\n  "currency": "USD",\n  "totalValue": 1200.00,\n  "positions": [...],\n  "availableCash": 150.00\n}'),
        ("GET /api/portfolio/core", "Returns all open CORE pool positions.", "[Position objects]"),
        ("GET /api/portfolio/satellite", "Returns all open SATELLITE pool positions.", "[Position objects]"),
        ("GET /api/fx", "Returns FX rate information (broker vs. market).",
         '{\n  "brokerRate": 1250.00,\n  "marketRate": 1245.00,\n  "spreadPct": 0.40\n}'),
        ("GET /api/settlement", "Returns settlement cash ledger for both markets.",
         '{\n  "ngx": { "availableCash": 850000.00, "settlingCash": 120000.00 },\n  "us": { "availableCash": 150.00, "settlingCash": 0.00 }\n}'),
        ("GET /api/performance", "Returns portfolio performance metrics.",
         '{\n  "snapshotDate": "2026-02-23",\n  "totalValue": 5250000.00,\n  "dailyPnl": 65000.00,\n  "dailyPnlPct": 1.25,\n  "unrealizedPnlByMarket": { "NGX": 45000, "US": 20000 },\n  "openPositionCount": 7\n}'),
        ("GET /api/dividends", "Returns all tracked dividend events.", "[DividendEvent objects]"),
    ]

    for endpoint, desc, response in endpoints:
        doc.add_heading(endpoint, level=2)
        doc.add_paragraph(desc)
        doc.add_paragraph("Response:")
        add_code_block(doc, response, "json")

    # ── Signals ──
    doc.add_heading("Signals & Intelligence", level=1)

    doc.add_heading("GET /api/signals", level=2)
    doc.add_paragraph("Returns today's trade signals. Optional filter: ?market=NGX or ?market=US")
    add_code_block(doc, '[\n  {\n    "symbol": "ZENITHBANK.XNSA",\n    "side": "BUY",\n    "strategy": "MomentumBreakout",\n    "confidenceScore": 78.5,\n    "entryPrice": 48.00,\n    "stopLoss": 45.00,\n    "targetPrice": 54.00\n  }\n]', "json")

    doc.add_heading("GET /api/news", level=2)
    doc.add_paragraph("Returns news items from the last 7 days.")

    doc.add_heading("GET /api/ai/cost", level=2)
    doc.add_paragraph("Returns AI usage cost tracking.")
    add_code_block(doc, '{\n  "date": "2026-02-23",\n  "dailyCostUsd": 1.25,\n  "monthlyCostUsd": 42.50,\n  "budgetExceeded": false\n}', "json")

    doc.add_heading("GET /api/ai/analysis/{symbol}", level=2)
    doc.add_paragraph("Returns AI analysis records for a stock from the last 30 days.")

    # ── Backtest ──
    doc.add_heading("Backtesting", level=1)

    doc.add_heading("POST /api/backtest/run", level=2)
    doc.add_paragraph("Triggers async backtest. Returns 202 Accepted.")
    doc.add_paragraph("Request body:")
    add_code_block(doc, '{\n  "strategyName": "MomentumBreakout",\n  "market": "NGX",\n  "startDate": "2025-02-01",\n  "endDate": "2026-02-01",\n  "initialCapital": 5000000.00\n}', "json")

    doc.add_heading("GET /api/backtest/runs", level=2)
    doc.add_paragraph("Returns all backtest runs. Optional filters: ?strategy=X&market=NGX")

    doc.add_heading("GET /api/backtest/runs/{id}", level=2)
    doc.add_paragraph("Returns a single backtest run by ID. 404 if not found.")

    doc.add_heading("GET /api/backtest/runs/{id}/trades", level=2)
    doc.add_paragraph("Returns simulated trades within a backtest run.")

    doc.add_heading("GET /api/backtest/runs/{id}/equity-curve", level=2)
    doc.add_paragraph("Returns equity curve data points for a backtest run.")

    doc.add_heading("GET /api/backtest/strategies", level=2)
    doc.add_paragraph("Returns all registered strategies with name, market, pool, and enabled status.")

    # ── System Control ──
    doc.add_heading("System Control", level=1)

    doc.add_heading("GET /api/killswitch", level=2)
    doc.add_paragraph("Returns kill switch status.")
    add_code_block(doc, '{\n  "active": false,\n  "reason": null,\n  "activatedAt": null\n}', "json")

    doc.add_heading("POST /api/killswitch/activate", level=2)
    doc.add_paragraph("Activates kill switch. Sends urgent WhatsApp + Telegram notification.")
    add_code_block(doc, 'Request:  { "reason": "Manual halt" }\nResponse: { "status": "activated", "reason": "Manual halt" }', "json")

    doc.add_heading("POST /api/killswitch/deactivate", level=2)
    doc.add_paragraph("Deactivates kill switch, allowing trading to resume.")

    doc.add_heading("POST /api/reconcile", level=2)
    doc.add_paragraph("Triggers manual reconciliation of positions and cash vs. broker.")
    add_code_block(doc, '{\n  "portfolioMatch": true,\n  "ngxCashMatch": true,\n  "usCashMatch": true,\n  "allClear": true\n}', "json")

    # ── Webhook ──
    doc.add_heading("Webhooks", level=1)

    doc.add_heading("POST /api/webhooks/whatsapp/message", level=2)
    doc.add_paragraph("Receives incoming WhatsApp messages from WAHA for trade approvals (YES/NO).")
    doc.add_paragraph("Always returns 200 OK. Non-approval messages are ignored.")

    # ── Health ──
    doc.add_heading("Health & Monitoring", level=1)
    add_styled_table(doc, ["Endpoint", "Description"], [
        ["GET /actuator/health", 'Application health (returns "UP")'],
        ["GET /actuator/info", "Application info"],
        ["GET /actuator/metrics", "Spring Boot metrics"],
        ["GET /actuator/flyway", "Flyway migration status"],
    ])

    # ── Request Flow ──
    doc.add_heading("Request Flow", level=1)
    add_diagram(doc, """sequenceDiagram
    participant UI as Dashboard UI
    participant API as REST API (8080)
    participant DB as PostgreSQL
    participant BOT as Trading Engine
    participant WAHA as WAHA (WhatsApp)

    UI->>API: GET /api/portfolio
    API->>DB: Query positions, snapshots
    DB-->>API: Position + snapshot data
    API-->>UI: Portfolio JSON

    UI->>API: POST /api/killswitch/activate
    API->>BOT: KillSwitchService.activate()
    BOT->>WAHA: Send urgent notification
    BOT->>DB: Persist kill switch state
    API-->>UI: activated

    WAHA->>API: POST /api/webhooks/whatsapp/message
    API->>BOT: TradeApprovalService.processReply()
    API-->>WAHA: 200 OK
    """, "Figure 2: API Request Flow", 5.5)

    save_doc(doc, "API_REFERENCE.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 7: DEPLOYMENT GUIDE
# ═══════════════════════════════════════════════════════════════════════════════

def generate_deployment_guide():
    print("\n[7/7] Generating DEPLOYMENT_GUIDE.docx...")
    doc = Document()
    setup_styles(doc)
    add_title_page(doc, "Deployment Guide", "Production Deployment & Operations",
                   "DevOps engineers and system administrators")

    # ── Architecture ──
    doc.add_heading("Architecture Overview", level=1)
    add_diagram(doc, """graph TB
    subgraph HOST["Production Host"]
        subgraph DOCKER["Docker Compose Stack"]
            BOT["trading-bot\\n:8080\\nJava 21 + Playwright"]
            PG["postgres\\n:5432\\nPostgreSQL 16"]
            WAHA["waha\\n:3000\\nWhatsApp WEBJS"]
        end
    end

    subgraph EXTERNAL["External Services"]
        EODHD["EODHD API"]
        TROVE["Trove Broker"]
        TELEGRAM["Telegram API"]
        ANTHROPIC["Anthropic API"]
        GMAIL["Gmail IMAP"]
    end

    BOT --> PG
    BOT --> WAHA
    BOT --> EODHD
    BOT --> TROVE
    BOT --> TELEGRAM
    BOT --> ANTHROPIC
    BOT --> GMAIL

    style BOT fill:#339af0,color:#fff
    style PG fill:#51cf66,color:#fff
    style WAHA fill:#ffa94d,color:#fff
    """, "Figure 1: Production Architecture", 5.5)

    # ── Prerequisites ──
    doc.add_heading("Prerequisites", level=1)
    add_styled_table(doc, ["Requirement", "Minimum", "Recommended"], [
        ["OS", "Ubuntu 22.04 / Debian 12", "Ubuntu 24.04 LTS"],
        ["Docker", "24.0+", "Latest stable"],
        ["Docker Compose", "v2.20+", "Latest stable"],
        ["RAM", "4 GB", "8 GB"],
        ["Disk", "20 GB", "50 GB"],
        ["CPU", "2 cores", "4 cores"],
    ])

    doc.add_heading("Required Outbound Access", level=2)
    add_styled_table(doc, ["Service", "Domain", "Port"], [
        ["EODHD", "eodhd.com", "443"],
        ["Trove Broker", "app.trovefinance.com", "443"],
        ["Telegram API", "api.telegram.org", "443"],
        ["Anthropic API", "api.anthropic.com", "443"],
        ["Gmail IMAP", "imap.gmail.com", "993"],
        ["WhatsApp Web", "web.whatsapp.com", "443"],
    ])

    # ── Environment Configuration ──
    doc.add_heading("Environment Configuration", level=1)
    doc.add_paragraph("Create a .env file in the project root. Never commit this file to version control.")
    add_code_block(doc, """# Database
DB_USERNAME=trader
DB_PASSWORD=<strong-password>

# Market Data
EODHD_API_KEY=<your-key>

# Broker
MERITRADE_USERNAME=<your-email>
MERITRADE_PASSWORD=<your-password>

# WhatsApp
WHATSAPP_CHAT_ID=<phone>@c.us
WAHA_API_KEY=<your-key>

# Telegram
TELEGRAM_BOT_TOKEN=<your-token>
TELEGRAM_CHAT_ID=<your-chat-id>

# AI (Optional)
AI_ENABLED=true
AI_API_KEY=<your-anthropic-key>""", "bash")

    # ── Deployment Steps ──
    doc.add_heading("Deployment Steps", level=1)

    doc.add_heading("Step 1: Clone and Configure", level=2)
    add_code_block(doc, "git clone <repository-url> trading-bot\ncd trading-bot\ncp .env.example .env\nnano .env", "bash")

    doc.add_heading("Step 2: Build and Start", level=2)
    add_code_block(doc, "docker compose up -d --build\ndocker compose ps", "bash")

    doc.add_heading("Step 3: Configure WhatsApp (WAHA)", level=2)
    doc.add_paragraph("1. Open http://<host>:3000 in a browser")
    doc.add_paragraph("2. Scan the QR code with WhatsApp mobile (Settings > Linked Devices)")
    doc.add_paragraph("3. Wait for session to establish (status: WORKING)")

    doc.add_heading("Step 4: Verify Health", level=2)
    add_code_block(doc, 'curl http://localhost:8080/actuator/health\n# Expected: {"status":"UP"}', "bash")

    # ── Startup Sequence ──
    doc.add_heading("Startup Sequence", level=1)
    add_diagram(doc, """sequenceDiagram
    participant DC as Docker Compose
    participant PG as PostgreSQL
    participant WAHA as WAHA
    participant BOT as Trading Bot

    DC->>PG: Start container
    PG->>PG: Initialize database
    PG->>DC: Healthcheck OK

    DC->>WAHA: Start container
    WAHA->>WAHA: Load WhatsApp session

    DC->>BOT: Start (depends_on postgres)
    BOT->>PG: Flyway: run 30 migrations
    PG-->>BOT: Migrations complete
    BOT->>BOT: Spring Boot init
    BOT->>WAHA: Verify connectivity
    BOT->>BOT: /actuator/health UP
    """, "Figure 2: Startup Sequence", 5.0)

    # ── Docker Services ──
    doc.add_heading("Docker Compose Services", level=1)

    doc.add_heading("trading-bot", level=2)
    add_styled_table(doc, ["Setting", "Value"], [
        ["Base image", "mcr.microsoft.com/playwright/java:v1.41.0-jammy"],
        ["JDK", "Eclipse Temurin 21"],
        ["Port", "8080"],
        ["Spring Profile", "prod"],
        ["JVM flags", "G1GC, heap dump on OOM"],
        ["User", "trader (non-root)"],
        ["Health check", "curl /actuator/health every 30s"],
    ])

    doc.add_heading("postgres", level=2)
    add_styled_table(doc, ["Setting", "Value"], [
        ["Image", "postgres:16-alpine"],
        ["Database", "ngx_trading"],
        ["Health check", "pg_isready every 10s"],
        ["Volume", "postgres_data"],
    ])

    doc.add_heading("waha", level=2)
    add_styled_table(doc, ["Setting", "Value"], [
        ["Image", "devlikeapro/waha:latest"],
        ["Engine", "WEBJS"],
        ["Port", "3000"],
        ["Volume", "waha_data (sessions)"],
    ])

    # ── Production Configuration ──
    doc.add_heading("Production Configuration", level=1)
    doc.add_paragraph("The prod profile (application-prod.yml) applies these overrides:")
    add_styled_table(doc, ["Setting", "Dev", "Prod", "Reason"], [
        ["meritrade.headless", "true", "true", "No GUI in production"],
        ["meritrade.slow-mo-ms", "500", "300", "Faster automation"],
        ["screenshot-retention-hours", "72", "48", "Less disk usage"],
        ["hikari.maximum-pool-size", "10", "15", "Handle concurrent load"],
        ["server.shutdown", "(default)", "graceful", "Clean shutdown"],
        ["logging.file.name", "(console)", "/app/logs/trading-bot.log", "File logging"],
        ["log max-file-size", "(default)", "50MB", "Prevents log bloat"],
        ["log max-history", "(default)", "30", "30-day retention"],
    ])

    # ── Monitoring ──
    doc.add_heading("Monitoring", level=1)

    doc.add_heading("Health Checks", level=2)
    add_code_block(doc, 'curl -s http://localhost:8080/actuator/health | jq .\ncurl -s http://localhost:8080/actuator/flyway | jq .', "bash")

    doc.add_heading("Key Log Patterns", level=2)
    add_styled_table(doc, ["Pattern", "Meaning", "Severity"], [
        ["KillSwitch ACTIVATED", "Trading halted", "CRITICAL"],
        ["Circuit breaker tripped", "Loss limit hit", "HIGH"],
        ["Order status: UNCERTAIN", "Execution failed", "HIGH"],
        ["OTP timeout", "OTP not received", "MEDIUM"],
        ["Budget exceeded", "AI spending limit", "LOW"],
        ["Scraper failed", "News source unreachable", "LOW"],
    ])

    doc.add_heading("Dashboard Queries", level=2)
    add_code_block(doc, """# Portfolio value
curl -s http://localhost:8080/api/portfolio | jq '.totalValueNgn'

# Kill switch status
curl -s http://localhost:8080/api/killswitch | jq '.'

# AI budget
curl -s http://localhost:8080/api/ai/cost | jq '.'""", "bash")

    # ── Backup & Recovery ──
    doc.add_heading("Backup & Recovery", level=1)

    doc.add_heading("Database Backup", level=2)
    add_code_block(doc, "# Manual backup\ndocker exec ngx-postgres pg_dump -U trader -d ngx_trading > backup.sql\n\n# Compressed backup\ndocker exec ngx-postgres pg_dump -U trader -d ngx_trading | gzip > backup.sql.gz", "bash")

    doc.add_heading("Automated Daily Backup (Cron)", level=2)
    add_code_block(doc, "# Add to crontab\n0 3 * * * docker exec ngx-postgres pg_dump -U trader -d ngx_trading | gzip > /backups/ngx_$(date +%Y%m%d).sql.gz\n0 4 * * * find /backups -name '*.sql.gz' -mtime +30 -delete", "bash")

    doc.add_heading("Restore", level=2)
    add_code_block(doc, "docker compose stop trading-bot\ncat backup.sql | docker exec -i ngx-postgres psql -U trader -d ngx_trading\ndocker compose start trading-bot", "bash")

    # ── Maintenance ──
    doc.add_heading("Maintenance", level=1)

    doc.add_heading("Updating the Application", level=2)
    add_code_block(doc, "git pull origin main\ndocker compose up -d --build trading-bot\ncurl -s http://localhost:8080/actuator/health | jq '.status'", "bash")
    add_callout(doc, "Flyway migrations run automatically on startup. New migrations apply when the updated container starts.")

    # ── Troubleshooting ──
    doc.add_heading("Troubleshooting", level=1)

    problems = [
        ("Bot Won't Start", "Check: docker logs ngx-trading-bot\nCommon causes: PostgreSQL not ready, missing .env vars, Flyway migration failure, port 8080 in use."),
        ("WhatsApp Not Working", "Check: curl http://localhost:3000/api/sessions\nFix: Restart WAHA and re-scan QR code."),
        ("OTP Failures", "Check: docker logs ngx-trading-bot | grep -i otp\nCommon: Gmail App Password expired, WAHA session disconnected."),
        ("Database Connection", "Check: docker exec ngx-postgres pg_isready -U trader -d ngx_trading"),
    ]
    for title, content in problems:
        doc.add_heading(title, level=2)
        add_code_block(doc, content, "")

    # ── Emergency Kill Switch ──
    doc.add_heading("Emergency Kill Switch", level=1)
    add_code_block(doc, """# Activate immediately
curl -X POST http://localhost:8080/api/killswitch/activate \\
  -H "Content-Type: application/json" \\
  -d '{"reason": "Emergency manual halt"}'

# Verify
curl -s http://localhost:8080/api/killswitch | jq .

# Deactivate when ready
curl -X POST http://localhost:8080/api/killswitch/deactivate""", "bash")

    # ── Security Checklist ──
    doc.add_heading("Security Checklist", level=1)
    checklist = [
        ".env file has restricted permissions (chmod 600 .env)",
        ".env is in .gitignore (never committed)",
        "PostgreSQL password is strong and unique",
        "WAHA API key is set (not default)",
        "Bot runs as non-root user (trader) inside container",
        "Ports 5432 and 3000 NOT exposed to public internet",
        "Port 8080 restricted to trusted networks or behind reverse proxy",
        "Gmail uses App Password (not main password)",
        "Telegram bot token not shared publicly",
        "Regular database backups configured and tested",
    ]
    for item in checklist:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)

    save_doc(doc, "DEPLOYMENT_GUIDE.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("NGX Trading Bot — Word Document Generator")
    print("=" * 60)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Output directory: {OUTPUT_DIR}")

    generate_developer_guide()
    generate_qa_guide()
    generate_business_overview()
    generate_product_spec()
    generate_pitch()
    generate_api_reference()
    generate_deployment_guide()

    print("\n" + "=" * 60)
    print("All 7 documents generated successfully!")
    print(f"Output: {OUTPUT_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()
